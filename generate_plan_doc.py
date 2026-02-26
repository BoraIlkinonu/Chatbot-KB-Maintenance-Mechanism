import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- Title Page ---
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Curriculum KB Maintenance Mechanism')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Strategy & Implementation Plan')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph('')

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run('17 February 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

classification = doc.add_paragraph()
classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = classification.add_run('Classification: Internal')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# --- Table of Contents ---
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Problem Statement',
    '2. Current State Analysis',
    '3. Proposed Architecture',
    '4. Repository Structure',
    '5. JSON Schema Design (Backward-Compatible)',
    '6. Conversion Pipeline',
    '7. Trigger Modes (Scheduled + Real-Time)',
    '8. Notification System',
    '9. Changelog & Diff Generation',
    '10. Team Workflow',
    '11. Implementation Phases',
    '12. Appendix: Sample JSON Output'
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()

# ===== SECTION 1 =====
doc.add_heading('1. Problem Statement', level=1)

doc.add_paragraph(
    'The Endstar AI Assistant chatbot serves students in the Explorer\'s Programme (G9\u2013G10, UAE). '
    'It draws knowledge from two separate pipelines:'
)

table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Pipeline', 'Owner', 'Content']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

table.rows[1].cells[0].text = 'Pipeline A (Our Team)'
table.rows[1].cells[1].text = 'Curriculum Team'
table.rows[1].cells[2].text = 'Term 1 & Term 2 curriculum data (lesson plans, rubrics, activities, assessment frameworks)'

table.rows[2].cells[0].text = 'Pipeline B (Partner Team)'
table.rows[2].cells[1].text = 'Endstar Platform Team'
table.rows[2].cells[2].text = 'Endstar technical data (tools, wiring, SDK, platform features, Wiki content)'

doc.add_paragraph('')
doc.add_heading('Identified Issues', level=2)

issues = [
    'Stale data: Pipeline A\u2019s KB was generated once from a CSV summary and never updated, even as source lesson plans evolved.',
    'Mixed content: Curriculum data was mixed with some Endstar technical data that should belong exclusively to Pipeline B.',
    'Shallow extraction: The CSV-to-JSON conversion captured metadata-level info but missed deep content \u2014 rubric table descriptors, detailed activity breakdowns, teacher notes from speaker notes, and assessment framework structures.',
    'False-positive QA flags: The QA assessment (11 Feb 2026) found that Hallucination Resistance scored 71.8% and Factual Accuracy scored 72.7%, with a significant portion of flags caused by KB gaps rather than actual chatbot defects.',
    'No update mechanism: There is no pipeline to regenerate the KB when source files are updated. Changes to lesson plans have no path to reach the chatbot.'
]
for issue in issues:
    doc.add_paragraph(issue, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Impact from QA Report (11 Feb 2026):')
run.bold = True

impact_table = doc.add_table(rows=6, cols=3)
impact_table.style = 'Light Grid Accent 1'
impact_table.alignment = WD_TABLE_ALIGNMENT.CENTER
impact_headers = ['Category', 'Score', 'Root Cause']
for i, h in enumerate(impact_headers):
    impact_table.rows[0].cells[i].text = h
    for paragraph in impact_table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

impact_data = [
    ['Term 1 Knowledge', '99.1%', 'KB adequate for Term 1'],
    ['Term 2 Knowledge', '75.7%', 'Missing rubric descriptors, incomplete assessment data'],
    ['Cross-Term Confusion', '69.6%', 'KB lacks structural awareness of term differences'],
    ['Student Simulation', '75.6%', 'Fabricated platform details due to missing KB content'],
    ['Hallucination Probing', '98.0%', 'Chatbot correctly defers when unsure'],
]
for i, row_data in enumerate(impact_data):
    for j, cell_data in enumerate(row_data):
        impact_table.rows[i+1].cells[j].text = cell_data

doc.add_page_break()

# ===== SECTION 2 =====
doc.add_heading('2. Current State Analysis', level=1)

doc.add_heading('Current JSON Structure', level=2)
doc.add_paragraph(
    'The existing KB file (Term 2 - Lesson Based Structure.json, 863KB) was generated from a CSV file. '
    'It contains 12 lessons with the following schema per lesson:'
)

schema_table = doc.add_table(rows=10, cols=3)
schema_table.style = 'Light Grid Accent 1'
schema_table.alignment = WD_TABLE_ALIGNMENT.CENTER
schema_headers = ['Field', 'Status', 'Issue']
for i, h in enumerate(schema_headers):
    schema_table.rows[0].cells[i].text = h
    for paragraph in schema_table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

schema_data = [
    ['lesson_title', 'Populated', 'Adequate'],
    ['metadata (16 sub-fields)', 'Populated', 'assessment_signals too shallow; endstar_tools sometimes incorrect'],
    ['metadata.images', 'Populated (66\u201378/lesson)', 'Well-structured with visual descriptions and tags'],
    ['description_of_activities', 'Populated', 'Thin \u2014 few sentences vs. full slide-by-slide breakdown'],
    ['videos_column', 'Empty', 'Never populated'],
    ['other_resources', 'Populated', 'Adequate'],
    ['testing_scores', 'Empty', 'Never populated'],
    ['comments', 'Empty', 'Never populated'],
    ['prompts', 'Empty', 'Never populated'],
]
for i, row_data in enumerate(schema_data):
    for j, cell_data in enumerate(row_data):
        schema_table.rows[i+1].cells[j].text = cell_data

doc.add_paragraph('')
doc.add_heading('What\u2019s Missing', level=2)
missing = [
    'key_facts: No key facts field exists in any lesson. The QA report flagged missing activities (Team Charter, team roles, timelines in Lesson 4).',
    'Rubric tables: assessment_signals has three one-line tiers (basic/intermediate/advanced) but no full rubric descriptors with criteria matrices.',
    'Teacher notes: Speaker notes from pptx files contain pedagogical guidance never captured.',
    'Assessment framework: Term-level assessment structure (portfolio 25%, product 50%, pitch 25%) and scoring methods not represented.',
    'Detailed activity breakdowns: Slide-by-slide activity sequences with timing and instructions.'
]
for m in missing:
    doc.add_paragraph(m, style='List Bullet')

doc.add_page_break()

# ===== SECTION 3 =====
doc.add_heading('3. Proposed Architecture', level=1)

doc.add_paragraph(
    'The mechanism consists of five components working together:'
)

arch_table = doc.add_table(rows=6, cols=3)
arch_table.style = 'Light Grid Accent 1'
arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
arch_headers = ['Component', 'Technology', 'Purpose']
for i, h in enumerate(arch_headers):
    arch_table.rows[0].cells[i].text = h
    for paragraph in arch_table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

arch_data = [
    ['Source Repository', 'GitHub (private repo) + GitHub Desktop', 'Version-controlled storage for pptx/docx source files with team collaboration'],
    ['Conversion Pipeline', 'Python (python-pptx, python-docx)', 'Extract content from source files preserving tables, speaker notes, slide structure'],
    ['Build Orchestrator', 'GitHub Actions', 'Trigger pipeline on schedule (midnight) or on push (real-time mode), controlled by config toggle'],
    ['Notification System', 'Slack Webhook', 'Alert distribution list when new KB build is ready, with changelog summary'],
    ['Output Artifacts', 'JSON files + changelog', 'Backward-compatible term KB JSONs with enriched nested data, plus human-readable diff'],
]
for i, row_data in enumerate(arch_data):
    for j, cell_data in enumerate(row_data):
        arch_table.rows[i+1].cells[j].text = cell_data

doc.add_paragraph('')

doc.add_heading('Pipeline Flow', level=2)
flow_steps = [
    'Team member updates a lesson plan (pptx/docx) in sources/term1/ or sources/term2/',
    'They commit and push via GitHub Desktop',
    'GitHub Actions detects the change (immediately in real-time mode, or at midnight in scheduled mode)',
    'The Python conversion script runs: extracts text, tables, speaker notes from all source files',
    'Script generates term1_kb.json and term2_kb.json with both original fields and new enriched block',
    'Script compares new JSON against previous version to generate a structured changelog',
    'Slack notification is sent to the distribution list with changelog summary and download link',
    'Chatbot maintainer retrieves updated JSONs from the repository'
]
for i, step in enumerate(flow_steps):
    doc.add_paragraph(f'{i+1}. {step}')

doc.add_page_break()

# ===== SECTION 4 =====
doc.add_heading('4. Repository Structure', level=1)

repo_lines = [
    'curriculum-kb/',
    '\u251c\u2500\u2500 sources/',
    '\u2502   \u251c\u2500\u2500 term1/                  \u2190 pptx, docx lesson plans (Term 1)',
    '\u2502   \u2514\u2500\u2500 term2/                  \u2190 pptx, docx lesson plans (Term 2)',
    '\u251c\u2500\u2500 output/',
    '\u2502   \u251c\u2500\u2500 term1_kb.json           \u2190 generated (do not edit manually)',
    '\u2502   \u251c\u2500\u2500 term2_kb.json           \u2190 generated (do not edit manually)',
    '\u2502   \u2514\u2500\u2500 changelog.md            \u2190 generated diff per build',
    '\u251c\u2500\u2500 scripts/',
    '\u2502   \u251c\u2500\u2500 convert.py              \u2190 main conversion pipeline',
    '\u2502   \u251c\u2500\u2500 diff_generator.py       \u2190 changelog/diff logic',
    '\u2502   \u2514\u2500\u2500 notify.py               \u2190 Slack notification sender',
    '\u251c\u2500\u2500 config.yaml                     \u2190 trigger mode, Slack webhook, dist list',
    '\u251c\u2500\u2500 .github/',
    '\u2502   \u2514\u2500\u2500 workflows/',
    '\u2502       \u2514\u2500\u2500 kb-build.yml         \u2190 GitHub Actions workflow',
    '\u251c\u2500\u2500 .gitignore',
    '\u2514\u2500\u2500 README.md'
]
for line in repo_lines:
    p = doc.add_paragraph(line)
    p.style = doc.styles['Normal']
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_paragraph('')
doc.add_heading('config.yaml Example', level=2)
config_lines = [
    'trigger_mode: scheduled          # Options: scheduled | realtime',
    'schedule_time: "0 20 * * *"      # Cron: midnight UAE (UTC+4 = 20:00 UTC)',
    '',
    'slack:',
    '  webhook_url: "https://hooks.slack.com/services/xxx/yyy/zzz"',
    '  channel: "#kb-updates"',
    '  notify:',
    '    - "@sarah"',
    '    - "@ahmed"',
    '    - "@curriculum-team"',
    '',
    'conversion:',
    '  extract_speaker_notes: true',
    '  extract_images: false           # Handled by existing pipeline',
    '  preserve_table_structure: true',
    '  output_dir: "output/"'
]
for line in config_lines:
    p = doc.add_paragraph(line)
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ===== SECTION 5 =====
doc.add_heading('5. JSON Schema Design (Backward-Compatible)', level=1)

doc.add_paragraph(
    'The key design principle is backward compatibility. All existing fields remain untouched. '
    'New rich data is added under a nested "enriched" object within each lesson. '
    'Systems reading the old fields continue to work. The chatbot maintainer can adopt the enriched '
    'data at their own pace.'
)

doc.add_heading('New Enriched Fields', level=2)

fields_table = doc.add_table(rows=8, cols=3)
fields_table.style = 'Light Grid Accent 1'
fields_table.alignment = WD_TABLE_ALIGNMENT.CENTER
f_headers = ['Field', 'Type', 'Description']
for i, h in enumerate(f_headers):
    fields_table.rows[0].cells[i].text = h
    for paragraph in fields_table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

f_data = [
    ['enriched.key_facts', 'Array<string>', 'Critical facts extracted from source files that the chatbot must know for accurate responses'],
    ['enriched.detailed_activities', 'Array<object>', 'Slide-by-slide activity breakdown with activity_id, title, description, and slide_references'],
    ['enriched.rubrics', 'Array<object>', 'Full rubric tables as structured data: title, headers[], rows[][] preserving exact row/column alignment'],
    ['enriched.teacher_notes', 'Array<string>', 'Pedagogical guidance from pptx speaker notes, prefixed with slide reference'],
    ['enriched.assessment_framework', 'Object', 'Term-level assessment structure: weights, scoring methods, rubric types'],
    ['enriched.source_files', 'Array<string>', 'List of source pptx/docx files that contributed to this lesson\u2019s enriched data'],
    ['enriched.last_updated', 'ISO 8601 string', 'Timestamp of when enriched data was last regenerated'],
]
for i, row_data in enumerate(f_data):
    for j, cell_data in enumerate(row_data):
        fields_table.rows[i+1].cells[j].text = cell_data

doc.add_paragraph('')
doc.add_heading('Rubric Table Structure (Critical)', level=2)
doc.add_paragraph(
    'Tables are the most important structural element to preserve. The QA report specifically flagged '
    'missing rubric descriptors as a major source of false-positive hallucination flags. '
    'Rubrics are stored as structured arrays to maintain exact row/column alignment:'
)

rubric_example = '''{
  "title": "Game Product Rubric (50%)",
  "headers": ["Criteria", "Approaching", "Meeting", "Above Expectations"],
  "rows": [
    ["Design coherence", "Basic layout with...", "Clear themed layout...", "Polished integrated..."],
    ["Mechanic implementation", "Simple mechanics...", "Working mechanics...", "Complex, balanced..."],
    ["User experience", "Functional but...", "Intuitive navigation...", "Seamless, engaging..."]
  ]
}'''
for line in rubric_example.split('\n'):
    p = doc.add_paragraph(line)
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_paragraph('')
doc.add_paragraph(
    'This structure guarantees that row[i][j] always corresponds to headers[j], '
    'eliminating any misalignment risk when the chatbot reads the data.'
)

doc.add_page_break()

# ===== SECTION 6 =====
doc.add_heading('6. Conversion Pipeline', level=1)

doc.add_paragraph(
    'The Python conversion script (scripts/convert.py) processes all source files and generates '
    'the output JSONs. It uses python-pptx and python-docx for extraction.'
)

doc.add_heading('Processing Steps', level=2)

proc_table = doc.add_table(rows=7, cols=3)
proc_table.style = 'Light Grid Accent 1'
proc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
p_headers = ['Step', 'Action', 'Details']
for i, h in enumerate(p_headers):
    proc_table.rows[0].cells[i].text = h
    for paragraph in proc_table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

p_data = [
    ['1. Discover', 'Scan source folders', 'List all .pptx and .docx files in sources/term1/ and sources/term2/'],
    ['2. Extract', 'Parse each file', 'For pptx: extract slide text, speaker notes, tables, slide order.\nFor docx: extract paragraphs, tables, headings, lists.'],
    ['3. Structure Tables', 'Preserve table layout', 'Convert each table to {headers: [], rows: [][]} format.\nHandle merged cells, empty cells, multi-line cell content.'],
    ['4. Map to Lessons', 'Associate content with lessons', 'Match source files to lesson numbers using filename conventions or folder structure.'],
    ['5. Build Enriched Block', 'Populate enriched fields', 'Assemble key_facts, detailed_activities, rubrics, teacher_notes, assessment_framework per lesson.'],
    ['6. Merge with Base', 'Combine with existing schema', 'Load existing base fields (from CSV or previous build), attach enriched block, output final JSON.'],
]
for i, row_data in enumerate(p_data):
    for j, cell_data in enumerate(row_data):
        proc_table.rows[i+1].cells[j].text = cell_data

doc.add_paragraph('')
doc.add_heading('Table Extraction Strategy', level=2)
doc.add_paragraph(
    'Tables in pptx and docx files require careful handling to preserve structure:'
)
table_strategies = [
    'Row-by-row extraction: Each table row becomes an array. Each cell becomes a string element within that array.',
    'Header detection: The first row is assumed to be headers unless explicitly marked otherwise.',
    'Merged cells: Merged cells are repeated in the array to maintain column alignment. A "merged" flag can be added if needed.',
    'Multi-line cells: Cell content with line breaks is preserved as-is within the string. No flattening.',
    'Empty cells: Represented as empty strings ("") to maintain positional integrity.',
    'Nested tables: Rare in lesson plans, but if encountered, the inner table is serialized as a sub-object within the cell.'
]
for s in table_strategies:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# ===== SECTION 7 =====
doc.add_heading('7. Trigger Modes (Scheduled + Real-Time)', level=1)

doc.add_paragraph(
    'The system supports two trigger modes, controlled by the trigger_mode field in config.yaml. '
    'Both modes use the same GitHub Actions workflow but with different trigger conditions.'
)

doc.add_heading('Scheduled Mode (Default)', level=2)
sched_points = [
    'Runs once daily at midnight UAE time (20:00 UTC)',
    'Checks if any source files changed since the last successful build',
    'If no changes detected, skips the build and does not send notifications',
    'Ideal for normal operations where multiple edits accumulate during the day',
    'Prevents unnecessary rebuilds from frequent small commits'
]
for s in sched_points:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Real-Time Mode', level=2)
rt_points = [
    'Triggers on every push to the main branch that includes changes in sources/ folder',
    'Useful for urgent KB fixes (e.g., correcting a rubric before a QA run)',
    'Activated by changing trigger_mode to "realtime" in config.yaml and pushing',
    'Can be toggled back to "scheduled" at any time',
    'Each push produces a separate build and notification'
]
for s in rt_points:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph('')
doc.add_heading('GitHub Actions Workflow Logic', level=2)
workflow_yaml = """name: KB Build Pipeline

on:
  schedule:
    - cron: "0 20 * * *"        # Midnight UAE
  push:
    branches: [main]
    paths: ["sources/**", "config.yaml"]

jobs:
  build:
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v4

      - name: Check trigger mode
        run: |
          MODE=$(python -c "import yaml; ...")
          if [[ "${{ github.event_name }}" == "push" \\
                && "$MODE" == "scheduled" ]]; then
            echo "Scheduled mode active, skipping push trigger"
            exit 0
          fi

      - name: Run conversion
        run: python scripts/convert.py

      - name: Generate changelog
        run: python scripts/diff_generator.py

      - name: Commit output
        run: |
          git add output/
          git commit -m "KB build: $(date -u)"
          git push

      - name: Notify Slack
        run: python scripts/notify.py"""

for line in workflow_yaml.split('\n'):
    p = doc.add_paragraph(line)
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Note: ')
run.bold = True
p.add_run(
    'GitHub Actions free tier provides 2,000 minutes/month for private repos. '
    'Each KB build takes approximately 1\u20133 minutes, well within budget even with daily scheduled runs.'
)

doc.add_page_break()

# ===== SECTION 8 =====
doc.add_heading('8. Notification System', level=1)

doc.add_paragraph(
    'Notifications are sent via Slack webhook when a new KB build completes successfully. '
    'The distribution list is managed in config.yaml so team members can be added or removed '
    'without modifying code.'
)

doc.add_heading('Slack Message Format', level=2)
slack_lines = [
    '\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501',
    '  KB Build Complete              #kb-updates',
    '\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501',
    '',
    '  New curriculum KB build ready.',
    '  Build time: 2026-02-17T20:00:12Z',
    '  Trigger: Scheduled (midnight UAE)',
    '',
    '  Changes:',
    '    Term 2 \u2014 Lesson 4:',
    '      + Added rubric descriptors (3 criteria x 3 levels)',
    '      + Added Team Charter activity details',
    '      ~ Updated assessment_signals',
    '    Term 2 \u2014 Lesson 7:',
    '      + Added detailed_activities (5 activities)',
    '      + Added teacher_notes (8 notes from speaker notes)',
    '',
    '  Files: term1_kb.json (no changes), term2_kb.json (updated)',
    '  Repo: https://github.com/your-org/curriculum-kb',
    '',
    '  cc: @sarah @ahmed @curriculum-team',
    '\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501',
]
for line in slack_lines:
    p = doc.add_paragraph(line)
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_paragraph('')
doc.add_heading('Notification Features', level=2)
notif_features = [
    'Build status: Success or failure with error details',
    'Trigger type: Whether the build was scheduled or real-time (push-triggered)',
    'Changelog summary: Which terms and lessons were affected, what was added/changed',
    'File status: Which output JSONs were updated vs. unchanged',
    'Repository link: Direct link to the repo for downloading updated files',
    'Distribution list: Configurable list of Slack users/groups to mention'
]
for n in notif_features:
    doc.add_paragraph(n, style='List Bullet')

doc.add_page_break()

# ===== SECTION 9 =====
doc.add_heading('9. Changelog & Diff Generation', level=1)

doc.add_paragraph(
    'Each build generates a structured changelog by comparing the new JSON output against the '
    'previous version. The diff is semantic, not a raw text diff \u2014 it describes what changed '
    'in human-readable terms.'
)

doc.add_heading('Diff Strategy', level=2)
diff_points = [
    'Lesson-level comparison: Detect added, removed, or modified lessons.',
    'Field-level comparison: For each lesson, compare enriched sub-fields (key_facts, rubrics, activities, etc.).',
    'Table-level comparison: For rubrics, detect added/removed rows, changed cell values, new columns.',
    'Summary statistics: Total files changed, lessons affected, fields added/modified/removed.',
    'Output format: Markdown file (changelog.md) committed to the output/ folder, plus a compact summary for Slack.'
]
for d in diff_points:
    doc.add_paragraph(d, style='List Bullet')

doc.add_heading('Example Changelog Entry', level=2)
changelog_example = """## KB Build \u2014 2026-02-17T20:00:12Z
**Trigger:** Scheduled (midnight UAE)
**Source commits since last build:** 3

### Term 2 Changes

#### Lesson 4 \u2013 Rewriting the Brief
- **enriched.key_facts:** Added 5 new facts (Team Charter, team roles, responsibilities, timelines, brief synthesis)
- **enriched.detailed_activities:** Added 3 activities (4.1 Team Formation, 4.2 Insight Wall, 4.3 Brief Rewriting)
- **enriched.rubrics:** Added Game Product Rubric (3 criteria \u00d7 3 levels)
- **enriched.teacher_notes:** Added 3 teacher notes from speaker notes

#### Lesson 7 \u2013 Prototype v2: Expanding Gameplay
- **enriched.detailed_activities:** Added 5 activities
- **enriched.teacher_notes:** Added 8 notes

### Term 1 Changes
No changes detected.

### Summary
| Metric | Value |
|--------|-------|
| Terms affected | 1 (Term 2) |
| Lessons modified | 2 |
| New key facts | 5 |
| New activities | 8 |
| New rubric tables | 1 |
| New teacher notes | 11 |"""

for line in changelog_example.split('\n'):
    p = doc.add_paragraph(line)
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ===== SECTION 10 =====
doc.add_heading('10. Team Workflow', level=1)

doc.add_paragraph(
    'The workflow is designed for team members who are comfortable with GitHub Desktop but not '
    'necessarily with command-line git or Python. All technical complexity is hidden behind the automation.'
)

doc.add_heading('For Curriculum Designers / Teachers', level=2)
teacher_steps = [
    'Open GitHub Desktop and pull the latest changes ("Fetch origin" then "Pull origin").',
    'Navigate to the sources/term1/ or sources/term2/ folder on your local machine.',
    'Edit the relevant pptx or docx file (e.g., update a rubric table, add activity details).',
    'Return to GitHub Desktop. You will see the changed files listed under "Changes".',
    'Write a brief commit message describing what you changed (e.g., "Updated Lesson 4 rubric descriptors").',
    'Click "Commit to main" then "Push origin".',
    'Done. The KB will rebuild automatically based on the current trigger mode.'
]
for i, step in enumerate(teacher_steps):
    doc.add_paragraph(f'{i+1}. {step}')

doc.add_heading('For KB Administrators', level=2)
admin_tasks = [
    'Toggle trigger mode: Edit config.yaml, change trigger_mode to "scheduled" or "realtime", commit and push.',
    'Add/remove notification recipients: Edit the slack.notify list in config.yaml.',
    'Monitor builds: Check the Actions tab in the GitHub repository for build logs and status.',
    'Manual trigger: Go to Actions tab > KB Build Pipeline > "Run workflow" button (for one-off builds).',
    'Review build history: Each build commits the output JSONs and changelog, creating a full audit trail in git.'
]
for a in admin_tasks:
    doc.add_paragraph(a, style='List Bullet')

doc.add_heading('For Chatbot Maintainer', level=2)
maintainer_tasks = [
    'Watch for Slack notifications in #kb-updates channel.',
    'Pull the latest output/term1_kb.json and output/term2_kb.json from the repository.',
    'Review output/changelog.md to understand what changed.',
    'Update the chatbot\u2019s curriculum KB endpoint with the new JSON files.',
    'Existing fields are unchanged \u2014 no integration changes needed unless adopting enriched data.'
]
for m in maintainer_tasks:
    doc.add_paragraph(m, style='List Bullet')

doc.add_page_break()

# ===== SECTION 11 =====
doc.add_heading('11. Implementation Phases', level=1)

phase_table = doc.add_table(rows=6, cols=4)
phase_table.style = 'Light Grid Accent 1'
phase_table.alignment = WD_TABLE_ALIGNMENT.CENTER
ph_headers = ['Phase', 'Deliverables', 'Dependencies', 'Priority']
for i, h in enumerate(ph_headers):
    phase_table.rows[0].cells[i].text = h
    for paragraph in phase_table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

ph_data = [
    ['Phase 1:\nRepository Setup', 'GitHub repo, folder structure, config.yaml, .gitignore, team access', 'GitHub account, team member GitHub usernames', 'Immediate'],
    ['Phase 2:\nConversion Pipeline', 'convert.py with pptx/docx extraction, table preservation, JSON generation', 'Source files committed to repo', 'High'],
    ['Phase 3:\nGitHub Actions', 'kb-build.yml with scheduled + real-time triggers, config toggle logic', 'Phase 2 complete', 'High'],
    ['Phase 4:\nDiff & Changelog', 'diff_generator.py, changelog.md generation with semantic comparison', 'Phase 2 complete', 'Medium'],
    ['Phase 5:\nSlack Notifications', 'notify.py, Slack webhook setup, distribution list management', 'Phase 3 + Phase 4 complete', 'Medium'],
]
for i, row_data in enumerate(ph_data):
    for j, cell_data in enumerate(row_data):
        phase_table.rows[i+1].cells[j].text = cell_data

doc.add_paragraph('')

doc.add_heading('Phase Dependencies', level=2)
doc.add_paragraph(
    'Phase 1 can begin immediately. Phase 2 requires the source files to be shared in the repository. '
    'Phases 3 and 4 can be developed in parallel once Phase 2 is functional. '
    'Phase 5 depends on both Phase 3 (for build triggers) and Phase 4 (for changelog content to include in notifications).'
)

doc.add_heading('Verification Plan', level=2)
verification = [
    'Phase 1: Clone repo via GitHub Desktop on a team member\u2019s machine, verify folder structure and push access.',
    'Phase 2: Run convert.py locally against source files. Compare output JSON with existing KB to verify backward compatibility. Validate that tables are correctly structured by checking row/column alignment.',
    'Phase 3: Push a test change to sources/ and verify GitHub Actions triggers correctly in both scheduled and real-time modes. Verify config toggle works.',
    'Phase 4: Make a known change to a source file, rebuild, and verify the changelog accurately reflects the change.',
    'Phase 5: Verify Slack notification arrives in the correct channel with correct changelog summary and mentions.',
    'End-to-end: Update a lesson plan pptx, push via GitHub Desktop, and verify the complete chain: build triggers \u2192 JSON regenerated \u2192 changelog generated \u2192 Slack notification sent.'
]
for v in verification:
    doc.add_paragraph(v, style='List Bullet')

doc.add_page_break()

# ===== SECTION 12 =====
doc.add_heading('12. Appendix: Sample JSON Output', level=1)

doc.add_paragraph('Below is a sample of what a single lesson would look like in the output JSON, '
                   'showing both the preserved original fields and the new enriched block:')

sample_json = """{
  "term": 2,
  "total_lessons": 12,
  "generated_from": "source files (pptx/docx)",
  "build_timestamp": "2026-02-17T20:00:12Z",
  "build_trigger": "scheduled",
  "lessons": [
    {
      "lesson_title": "Lesson 4 \u2013 Rewriting the Brief",
      "url": "Lesson 4",
      "metadata": {
        "term_id": 2,
        "lesson_id": 4,
        "title": "Rewriting the Brief: From Insights to Action",
        "grade_band": "G9\u2013G10",
        "core_topics": ["Brief rewriting", "Team formation"],
        "learning_objectives": ["..."],
        "assessment_signals": ["basic: ...", "intermediate: ...", "advanced: ..."],
        "images": ["... (existing image data preserved) ..."]
      },
      "description_of_activities": "Students rewrite their design brief...",
      "other_resources": "Brief Template, Team Charter Template",
      "videos_column": "",
      "testing_scores": "",
      "comments": "",
      "prompts": "",

      "enriched": {
        "key_facts": [
          "Students form teams of 3-4 and assign roles",
          "Roles: Project Manager, Lead Designer, Developer, QA Tester",
          "Team Charter defines responsibilities and deadlines",
          "Brief rewriting synthesises insights from Lessons 1-3",
          "Output: Revised brief with problem, audience, constraints"
        ],
        "detailed_activities": [
          {
            "activity_id": "4.1",
            "title": "Team Formation & Charter",
            "description": "Students form teams, assign roles...",
            "slide_references": [3, 4, 5]
          },
          {
            "activity_id": "4.2",
            "title": "Insight Wall",
            "description": "Teams compile research findings...",
            "slide_references": [7, 8]
          },
          {
            "activity_id": "4.3",
            "title": "Brief Rewriting",
            "description": "Using insights, teams rewrite their brief...",
            "slide_references": [10, 11, 12]
          }
        ],
        "rubrics": [
          {
            "title": "Game Product Rubric (50%)",
            "headers": ["Criteria", "Approaching", "Meeting",
                        "Above Expectations"],
            "rows": [
              ["Design coherence",
               "Basic layout with minimal theming",
               "Clear themed layout with consistent visual language",
               "Polished, integrated design with strong identity"],
              ["Mechanic implementation",
               "Simple mechanics with limited interaction",
               "Working mechanics that support gameplay loop",
               "Complex, balanced mechanics with emergent gameplay"]
            ]
          }
        ],
        "teacher_notes": [
          "Slide 3: Allow 10 mins for team formation. Max 4/team.",
          "Slide 5: Students often skip Charter. Emphasise portfolio.",
          "Slide 10: Brief must address ALL five elements."
        ],
        "assessment_framework": {
          "portfolio_weight": "25%",
          "product_weight": "50%",
          "pitch_weight": "25%",
          "portfolio_scoring": "complete / partial / missing",
          "product_rubric_type": "Approaching / Meeting / Above Expectations"
        },
        "source_files": [
          "Lesson 4 - Rewriting the Brief.pptx",
          "Teacher Assessment Guide.docx"
        ],
        "last_updated": "2026-02-17T20:00:12Z"
      }
    }
  ]
}"""

for line in sample_json.split('\n'):
    p = doc.add_paragraph(line)
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

# --- Footer note ---
doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('--- End of Document ---')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(9)

# --- Save ---
output_path = 'KB_Maintenance_Mechanism_Plan.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
