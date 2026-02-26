"""
Generate comprehensive pipeline documentation v3 as DOCX.
Updated with: existing pipeline context, smart change analyzer,
validation post-build check, removal of non-GitHub stages.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def code(text, size=8):
    for line in text.split('\n') if isinstance(text, str) else text:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = 'Consolas'
            r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    doc.add_paragraph('')

def tbl(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            t.rows[i+1].cells[j].text = str(cell)
            for p in t.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph('')

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def bold_para(bold_text, normal_text):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    p.add_run(normal_text)

# ============================================================
# TITLE
# ============================================================
for _ in range(5):
    doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Curriculum KB Maintenance Mechanism')
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Comprehensive Pipeline Design & Technical Specification')
r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('25 February 2026 | Version 3.0')
r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Classification: Internal | Incorporates Existing Pipeline Analysis')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# TOC
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc = [
    '1. Executive Summary',
    '2. Problem Statement & QA Impact',
    '3. Design Principles & Key Decisions',
    '4. Source Folder Analysis (Live Scan — 146 Files, 3 Terms)',
    '5. Existing Pipeline Analysis (Term 2 Reference)',
    '6. Complete Pipeline Architecture',
    '7. Google Drive Integration Layer',
    '   7.1 APIs & Permissions (6 APIs)',
    '   7.2 OAuth Authentication',
    '   7.3 Change Detection',
    '   7.4 File Format Handling (Native vs Office)',
    '8. Sync Layer — Detailed Design',
    '   8.1 Trigger Modes',
    '   8.2 Step-by-Step Sync Process',
    '   8.3 Validation Checks',
    '   8.4 File Scenarios',
    '9. Comprehensive Logging System',
    '   9.1 Log Everything — No Filtering',
    '   9.2 Log Entry Schema (Exhaustive)',
    '   9.3 Log Security (Read-Only)',
    '   9.4 Activity Tracking Capabilities & Limitations',
    '10. Conversion Pipeline (Smart Multi-Stage)',
    '   10.1 Change Analyzer — What Stages to Re-run',
    '   10.2 Stage 1: Media Extraction',
    '   10.3 Stage 2: Document Conversion',
    '   10.4 Stage 3: Native Google Content Extraction',
    '   10.5 Stage 4: Image Analysis (Manual/Admin — Claude)',
    '   10.6 Stage 5: Consolidation',
    '   10.7 Stage 6: KB Building',
    '   10.8 Stage 7: Automated Validation (Post-Build Check)',
    '   10.9 Table Preservation Strategy',
    '   10.10 Content-Level Diff',
    '11. JSON Schema Design (Backward-Compatible)',
    '12. Notification System (Slack)',
    '13. GitHub Repository Design',
    '14. Conflict & Error Handling',
    '15. Team Workflow (Zero Extra Steps)',
    '16. Testing & Simulation Plan',
    '17. Implementation Phases',
    '18. Appendix A: config.yaml Reference',
    '19. Appendix B: Full Log Entry Schema',
    '20. Appendix C: Validation Signal Definitions',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This document defines the complete design for an automated Curriculum Knowledge Base (KB) '
    'maintenance mechanism. The system monitors Google Drive folders containing curriculum source '
    'files (pptx, docx, Google Slides, Google Docs), detects changes, validates them, converts '
    'content to structured JSON, runs automated quality validation, and notifies the team via Slack '
    '\u2014 all without requiring any change to how the curriculum team currently works.'
)
doc.add_paragraph(
    'The design incorporates learnings from the existing Term 2 pipeline (a 7-stage, 16-script '
    'system that produced 61 markdown files, 443 AI-described images, and 3 video transcripts). '
    'Non-GitHub-compatible stages (ffmpeg, Whisper transcription) are excluded. The pipeline '
    'intelligently detects what type of change occurred and only re-runs the necessary stages.'
)
bold_para('Core Principle: ', 'Users edit files on Google Drive exactly as they do today. No Git, no GitHub Desktop, no extra steps, no locking, no restrictions. The entire pipeline is invisible to content creators.')

doc.add_page_break()

# ============================================================
# 2. PROBLEM STATEMENT
# ============================================================
doc.add_heading('2. Problem Statement & QA Impact', level=1)
doc.add_heading('Current Problems', level=2)
for p in [
    'Stale KB data: Generated once from a CSV summary, never updated as source lesson plans evolved.',
    'Mixed content: Curriculum data mixed with Endstar technical data belonging to the other team\u2019s pipeline.',
    'Shallow extraction: CSV-to-JSON captured metadata only \u2014 missed rubric descriptors, activity breakdowns, teacher notes, assessment frameworks.',
    'No update mechanism: Changes to lesson plans on Google Drive have no path to reach the chatbot KB.',
    'No audit trail: No record of who changed what, when, or why.',
    'No automated validation: The existing 5-signal validation system was designed but never executed.',
]:
    bullet(p)

doc.add_heading('QA Impact (11 February 2026)', level=2)
tbl(['Category', 'Score', 'Root Cause'], [
    ['Term 1 Knowledge', '99.1%', 'KB adequate for Term 1'],
    ['Term 2 Knowledge', '75.7%', 'Missing rubric descriptors, incomplete assessment data'],
    ['Cross-Term Confusion', '69.6%', 'KB lacks structural awareness of term differences'],
    ['Student Simulation', '75.6%', 'Fabricated platform details due to missing KB content'],
    ['Hallucination Probing', '98.0%', 'Chatbot correctly defers when unsure'],
    ['Overall Factual Accuracy', '72.7%', 'KB gaps, not chatbot defects'],
    ['Overall Hallucination Resistance', '71.8%', 'False positives from incomplete ground truth'],
])
doc.add_page_break()

# ============================================================
# 3. DESIGN PRINCIPLES
# ============================================================
doc.add_heading('3. Design Principles & Key Decisions', level=1)
for title, desc in [
    ('Google Drive is the source of truth', 'Users work on Google Drive. They do not interact with Git, GitHub, or any new tool.'),
    ('No file locking or ownership restrictions', 'Anyone can edit any file at any time. The system detects and handles issues after the fact.'),
    ('Log everything, restrict nothing', 'Every file, every sync, every action logged exhaustively. UNCHANGED files logged as unchanged. Append-only, read-only protected.'),
    ('Backward-compatible JSON schema', 'Existing fields preserved exactly. New data in nested "enriched" block.'),
    ('Scheduled + on-demand triggers', 'Default: midnight UAE daily. Admin can trigger on-demand.'),
    ('Smart stage re-execution', 'The pipeline detects what type of change occurred and only re-runs necessary stages, not the entire 7-stage pipeline.'),
    ('Automated validation post-build', 'Every KB build is validated using a 5-signal consensus system. Anomalies flagged in Slack. Critical issues block publishing.'),
    ('No non-Python dependencies', 'ffmpeg, Whisper, and other system-level tools excluded. Pipeline runs on standard GitHub Actions Ubuntu runner with Python packages only.'),
    ('Slack for all notifications', 'Activity alerts, build notifications, validation results, error warnings all go to Slack.'),
    ('Zero user workflow changes', 'Content creators\u2019 workflow: open file, edit, save. Nothing else.'),
]:
    bold_para(f'{title}: ', desc)
doc.add_page_break()

# ============================================================
# 4. SOURCE FOLDER ANALYSIS
# ============================================================
doc.add_heading('4. Source Folder Analysis (Live Scan \u2014 146 Files, 3 Terms)', level=1)
doc.add_paragraph('Based on live Google Drive API scan performed 25 February 2026.')

doc.add_heading('Overview', level=2)
tbl(['Folder', 'Content', 'Files', 'Subfolders', 'Size'], [
    ['Term 1 (Foundations)', 'Jan 2026 Cohort', '51', '14', '481 MB'],
    ['Term 2 (Accelerator)', 'Aug 2025 Cohort', '75', '45', '1,069 MB'],
    ['Term 3 (Mastery)', 'Aug 2025 Cohort', '20', '15', '510 MB'],
    ['TOTAL', '', '146', '74', '2,060 MB'],
])

doc.add_heading('File Type Distribution', level=2)
tbl(['File Type', 'Count', 'Size', 'Pipeline Action'], [
    ['PowerPoint (.pptx)', '~50', '~800MB', 'Extract text, tables, speaker notes via python-pptx'],
    ['Google Slides (native)', '~8', 'N/A (cloud)', 'Extract via Slides API (structured JSON)'],
    ['Word (.docx)', '~15', '~5MB', 'Extract via python-docx'],
    ['Google Docs (native)', '~8', 'N/A (cloud)', 'Extract via Docs API (structured JSON)'],
    ['Google Sheets (native)', '~3', 'N/A (cloud)', 'Extract via Sheets API (cell data)'],
    ['PDF', '~25', '~170MB', 'Store only. Not parsed for KB.'],
    ['Video (.mp4, .mov)', '~10', '~750MB', 'Store only. Not processed (no ffmpeg/Whisper on GitHub).'],
    ['Excel (.xlsx)', '1', '46KB', 'Parse via openpyxl'],
])

doc.add_heading('Naming Inconsistencies Detected', level=2)
tbl(['Issue', 'Example', 'Impact'], [
    ['Trailing spaces', '"Lesson 1 .pptx"', 'Fuzzy match normalizes whitespace'],
    ['Typos', '"Exampler Work"', 'Log and accept'],
    ['Singular vs plural folders', '"Lesson Plans" vs "Lesson Plan"', 'Flexible folder scanning'],
    ['Inconsistent separators', '"Lesson 1" vs "Lesson1"', 'Normalization in duplicate check'],
    ['Missing titles', 'Term 3: "Lesson 5.pptx" vs Term 1: "Explorer\'s Programme - Lesson 5.pptx"', 'Lesson number extraction handles both'],
])

doc.add_heading('Key Structural Differences Between Terms', level=2)
tbl(['Aspect', 'Term 1', 'Term 2', 'Term 3'], [
    ['Lessons', '22', '12', '14'],
    ['Teacher Slides', 'Yes (single set)', 'Yes (separate from student)', 'Yes'],
    ['Student Slides', 'Same as teacher', 'Separate (often Google native)', 'No'],
    ['Lesson Plan docs', 'No', 'Yes (12 docx files)', 'No'],
    ['Assessment guides', 'PDF + Google Docs', 'docx + PDF', 'Empty (in development)'],
    ['Exemplar work', 'PDF workbooks', 'pptx per week + videos', 'No'],
    ['Videos', '1 (98MB, tech support)', '3 (curriculum-related)', '4 (curriculum) + 2 (resources)'],
    ['Native Google files', 'Some', 'Many (Student Slides, Sheets)', 'None detected'],
])
doc.add_page_break()

# ============================================================
# 5. EXISTING PIPELINE ANALYSIS
# ============================================================
doc.add_heading('5. Existing Pipeline Analysis (Term 2 Reference)', level=1)
doc.add_paragraph(
    'The Term 2 KB was built using a 7-stage, 16-script pipeline. This section documents what '
    'that pipeline does, which parts we can reuse, and which are excluded from the new mechanism.'
)

doc.add_heading('Original 7-Stage Pipeline', level=2)
tbl(['Stage', 'Script(s)', 'Purpose', 'New Pipeline Status'], [
    ['1. Media Extraction', 'extract_media.py', 'Extract images from PPTX (ZIP/XML parsing), video keyframes via ffmpeg, audio via ffmpeg', 'PARTIAL \u2014 Image extraction reusable. ffmpeg parts EXCLUDED.'],
    ['2. Document Conversion', 'convert_docs.py, verify_conversions.py', 'Convert DOCX/PPTX/PDF/XLSX to Markdown. Verify text preservation.', 'REUSABLE \u2014 Core of new pipeline. Runs on GitHub Actions.'],
    ['3. Video Transcription', 'transcribe_videos.py', 'Whisper speech-to-text on extracted audio', 'EXCLUDED \u2014 Requires Whisper/torch. Not available on GitHub Actions.'],
    ['4. Image Analysis', 'vision_claude_pipeline.py + Claude Code agents', 'AI-generated image descriptions (443 images described)', 'MANUAL \u2014 Requires Claude multimodal. Admin triggered. Only for NEW images.'],
    ['5. Consolidation', 'merge_slide_metadata.py + inline script', 'Merge all batch results into MASTER file with slide metadata', 'REUSABLE \u2014 Runs on GitHub Actions.'],
    ['6. KB Building', 'compare_and_build_kb.py, fix_csv_and_add_images.py', 'Build structured JSON KB from all sources', 'REUSABLE \u2014 Refactored for new enriched{} schema.'],
    ['7. Validation', 'validation_parser.py, validation_mapper.py, validation_anomalies.py, validation_report.py', '5-signal consensus validation (designed but never executed)', 'INCLUDED \u2014 Automated post-build check. Runs every build.'],
])

doc.add_heading('What We Keep vs Exclude', level=2)
tbl(['Component', 'Status', 'Reason'], [
    ['python-pptx parsing', 'KEEP', 'Pure Python, runs on GitHub Actions'],
    ['python-docx parsing', 'KEEP', 'Pure Python, runs on GitHub Actions'],
    ['openpyxl (Excel)', 'KEEP', 'Pure Python, runs on GitHub Actions'],
    ['PyPDF2 (PDF text)', 'KEEP', 'Pure Python, store text but not for KB'],
    ['Google Slides API', 'NEW', 'For native Google Slides content extraction'],
    ['Google Docs API', 'NEW', 'For native Google Docs content extraction'],
    ['Google Sheets API', 'NEW', 'For native Google Sheets content extraction'],
    ['ffmpeg (keyframes)', 'EXCLUDE', 'System dependency, not available on GitHub Actions free tier reliably'],
    ['Whisper (transcription)', 'EXCLUDE', 'Requires torch + GPU. Too heavy for GitHub Actions.'],
    ['Gemini Vision API', 'EXCLUDE', 'Failed in original pipeline (SDK issues). Not needed.'],
    ['Claude Code agents (image desc)', 'MANUAL', 'Only triggered when new images detected. Admin runs separately.'],
    ['Validation scripts (4)', 'KEEP', 'Pure Python. Now automated as post-build check.'],
])

doc.add_heading('Original Pipeline Statistics (Term 2)', level=2)
tbl(['Metric', 'Value'], [
    ['Source files processed', '51'],
    ['Markdown files generated', '61'],
    ['Images extracted from PPTX', '~750'],
    ['Images after AI filtering', '443'],
    ['Video keyframes extracted', '72'],
    ['Video transcripts', '3'],
    ['Lessons with full coverage', '12/12'],
    ['Python scripts', '16'],
])
doc.add_page_break()

# ============================================================
# 6. COMPLETE PIPELINE ARCHITECTURE
# ============================================================
doc.add_heading('6. Complete Pipeline Architecture', level=1)
code("""\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502  GOOGLE DRIVE (Source of Truth)                              \u2502
\u2502  Term 1 (22 lessons) | Term 2 (12 lessons) | Term 3 (14)    \u2502
\u2502  Anyone edits. No restrictions. No extra steps.              \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
               \u2502  Drive API + Slides/Docs/Sheets APIs (6 scopes)
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502  SYNC LAYER (GitHub Actions cloud runner)                    \u2502
\u2502  1. Detect changes  2. Log everything  3. Validate           \u2502
\u2502  4. Download  5. Commit to Git  6. Notify Slack              \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
               \u2502
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502  CHANGE ANALYZER                                             \u2502
\u2502  Compares new files against previous sync to determine:      \u2502
\u2502  \u2022 Text/table/notes only? \u2192 Stages 2, 6, 7                  \u2502
\u2502  \u2022 New images detected? \u2192 Stages 1, 2, 5, 6, 7 + flag admin \u2502
\u2502  \u2022 Images removed? \u2192 Stages 1, 5, 6, 7                      \u2502
\u2502  \u2022 New lesson file? \u2192 All stages + flag admin for images     \u2502
\u2502  \u2022 Lesson plan docx only? \u2192 Stages 2, 6, 7                  \u2502
\u2502  \u2022 Native Google file? \u2192 Stage 3 (API), 6, 7                \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
               \u2502
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502  CONVERSION PIPELINE (7 Stages \u2014 selective execution)       \u2502
\u2502                                                              \u2502
\u2502  Stage 1: Media Extraction (python-pptx ZIP/XML)             \u2502
\u2502  Stage 2: Document Conversion (pptx/docx \u2192 markdown/JSON)  \u2502
\u2502  Stage 3: Native Google Extraction (Slides/Docs/Sheets API)  \u2502
\u2502  Stage 4: Image Analysis (MANUAL \u2014 Claude, admin only)      \u2502
\u2502  Stage 5: Consolidation (merge all extracted content)        \u2502
\u2502  Stage 6: KB Building (structured JSON with enriched{})      \u2502
\u2502  Stage 7: Validation (5-signal post-build check)             \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
               \u2502
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502  OUTPUTS & NOTIFICATIONS                                    \u2502
\u2502  JSON: term1_kb.json, term2_kb.json, term3_kb.json           \u2502
\u2502  Changelog: Content-level diff (cell-level for tables)       \u2502
\u2502  Validation: Pass/fail per lesson, anomaly list              \u2502
\u2502  Logs: Complete audit trail                                  \u2502
\u2502  Slack: Sync alerts, build results, validation results       \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518""")
doc.add_page_break()

# ============================================================
# 7. GOOGLE DRIVE INTEGRATION
# ============================================================
doc.add_heading('7. Google Drive Integration Layer', level=1)

doc.add_heading('7.1 APIs & Permissions', level=2)
tbl(['#', 'API', 'Scope', 'Purpose', 'Admin?'], [
    ['1', 'Drive API v3', 'drive.readonly', 'File listing, download, metadata, revisions, permissions, checksums', 'No'],
    ['2', 'Drive Activity API v2', 'drive.activity.readonly', 'Action feed: edit/rename/move/delete/comment/share events', 'No'],
    ['3', 'Slides API v1', 'presentations.readonly', 'Structured content from Google Slides (text, tables, notes)', 'No'],
    ['4', 'Docs API v1', 'documents.readonly', 'Structured content from Google Docs (paragraphs, tables)', 'No'],
    ['5', 'Sheets API v4', 'spreadsheets.readonly', 'Cell data from Google Sheets', 'No'],
    ['6', 'People API', 'directory.readonly', 'Resolve user IDs to profiles (name, email, photo)', 'No'],
    ['7', 'Admin SDK (DEFERRED)', 'admin.reports.audit.readonly', 'View/download events \u2014 requires Workspace admin', 'YES'],
])

doc.add_heading('7.2 OAuth Authentication', level=2)
for s in [
    'Google Cloud Project: poetic-dock-483707-n1',
    'Client type: Installed application (Desktop)',
    'First run: Opens browser for Google consent screen',
    'Token stored as GitHub Actions encrypted secret (GOOGLE_TOKEN_JSON)',
    'Auto-refresh using refresh token \u2014 no re-login needed',
    'All 6 non-admin scopes requested at initial auth',
]:
    bullet(s)

doc.add_heading('7.3 Change Detection', level=2)
doc.add_paragraph('The sync script maintains sync_state.json recording the last-known state of every file.')
code("""Comparison logic:
  File in Drive + File in sync_state + same md5      \u2192 UNCHANGED (logged)
  File in Drive + File in sync_state + different md5  \u2192 MODIFIED
  File in Drive + NOT in sync_state                   \u2192 ADDED
  NOT in Drive  + File in sync_state                  \u2192 DELETED""")

doc.add_heading('7.4 File Format Handling', level=2)
code("""For each file detected:
\u251c\u2500 Google Slides (native)  \u2192 Slides API \u2192 structured JSON
\u251c\u2500 Google Docs (native)    \u2192 Docs API \u2192 structured JSON
\u251c\u2500 Google Sheets (native)  \u2192 Sheets API \u2192 cell data
\u251c\u2500 PowerPoint (.pptx)      \u2192 Download \u2192 python-pptx parse
\u251c\u2500 Word (.docx)            \u2192 Download \u2192 python-docx parse
\u251c\u2500 Excel (.xlsx)           \u2192 Download \u2192 openpyxl parse
\u251c\u2500 PDF                     \u2192 Download \u2192 Store only, not parsed for KB
\u2514\u2500 Video (.mp4/.mov)       \u2192 Log metadata only, not processed

Log note per file: "content_extraction_method": "<method used>" """)
doc.add_page_break()

# ============================================================
# 8. SYNC LAYER
# ============================================================
doc.add_heading('8. Sync Layer \u2014 Detailed Design', level=1)

doc.add_heading('8.1 Trigger Modes', level=2)
tbl(['Mode', 'Trigger', 'Who', 'When'], [
    ['Scheduled', 'Cron: midnight UAE (20:00 UTC)', 'Automatic', 'Daily \u2014 picks up all changes from the day'],
    ['On-demand', 'GitHub Actions "Run workflow"', 'Admin only', 'Urgent updates, testing, ad-hoc sync'],
])

doc.add_heading('8.2 Step-by-Step Sync Process', level=2)
for title, desc in [
    ('Step 1 \u2014 DETECT: ', 'For each tracked folder in config.yaml, call Drive API. Compare against sync_state.json. Categorize every file as UNCHANGED, MODIFIED, ADDED, or DELETED. Scan parent folder for untracked items.'),
    ('Step 2 \u2014 LOG: ', 'For EVERY file (including UNCHANGED), capture complete API data: user, file metadata, revision history, permissions, capabilities, checksums, links, Drive Activity actions. Append to logs/activity_log.json (append-only, protected).'),
    ('Step 3 \u2014 VALIDATE: ', 'Run checks on changed files: duplicate detection (3 layers: exact, fuzzy, content), deletion protection, integrity parsing, lesson count guard. Each file marked ACCEPT or HOLD.'),
    ('Step 4 \u2014 DOWNLOAD: ', 'Download ACCEPTED files to GitHub Actions cloud runner (NOT local machine). For native Google files, also extract via APIs. Place in sources/termN/.'),
    ('Step 5 \u2014 COMMIT: ', 'Stage sources/, logs/, sync_state.json. Commit with descriptive message. Push. If nothing changed, skip.'),
    ('Step 6 \u2014 NOTIFY: ', 'Slack message: synced files, who changed them, warnings, held files. Triggers conversion pipeline.'),
]:
    bold_para(title, desc)

doc.add_heading('8.3 Validation Checks', level=2)
for title, desc in [
    ('DUPLICATE CHECK (3 layers): ', 'Layer 1: Detect "(1)" copies from Drive\'s "Keep both". Layer 2: Fuzzy name match (normalize, extract lesson number, Levenshtein distance > 0.8). Layer 3: MD5 content comparison.'),
    ('DELETION CHECK: ', 'File missing from Drive? Keep in Git, alert admin. Google Drive keeps deleted files 30 days.'),
    ('INTEGRITY CHECK: ', 'Attempt python-pptx/docx parse. Corrupt file? Reject it, continue others, alert.'),
    ('LESSON COUNT CHECK: ', 'Fewer lessons than previous sync? Hold entire term build, alert admin.'),
]:
    bold_para(title, desc)

doc.add_heading('8.4 File Scenarios', level=2)
for title, desc in [
    ('Normal edit: ', 'User edits and saves. Drive creates new revision. Sync detects via md5 change. Downloads latest. Commits. Logs.'),
    ('Upload with "Replace": ', 'New revision of existing file. Same as normal edit.'),
    ('Upload with "Keep both": ', 'Creates "Filename (1).pptx". Duplicate check Layer 1 catches it. HELD.'),
    ('Similar name (typo): ', 'Fuzzy match catches "Lesson 4.pptx" vs "Lesson 4 - Rewriting the Brief.pptx". HELD.'),
    ('File deleted: ', 'KB keeps previous version. Alert admin for confirmation.'),
    ('New file (new lesson): ', 'Accepted if naming valid. Slack notification.'),
    ('New term folder: ', 'Not auto-tracked. Alert: "Add to config.yaml to include."'),
    ('New subfolder in tracked term: ', 'Auto-scanned. Slack notification.'),
    ('Unrelated folder/file in parent: ', 'Logged and reported. Not synced.'),
    ('Corrupt file: ', 'Rejected. Other files continue. Alert.'),
]:
    bold_para(title, desc)
doc.add_page_break()

# ============================================================
# 9. LOGGING
# ============================================================
doc.add_heading('9. Comprehensive Logging System', level=1)

doc.add_heading('9.1 Log Everything \u2014 No Filtering', level=2)
doc.add_paragraph(
    'The logging system captures EVERY piece of data returned by every API call for every file '
    'on every sync cycle. There is no filtering. If the API returns it, we store it. '
    'UNCHANGED files are logged as unchanged with their full current state. The log is a '
    'complete snapshot of the entire Drive state at every sync point.'
)

doc.add_heading('9.2 Log Entry Schema', level=2)
doc.add_paragraph('Each sync run produces one entry. Key data per file:')
tbl(['Category', 'Fields', 'Source'], [
    ['Sync run', 'sync_id, trigger, timestamps, duration, GitHub run ID/URL, totals per status', 'System'],
    ['Folder scan', 'folder name/ID/link, file count, subfolder list, total size', 'Drive API'],
    ['Parent activity', 'New folders/files, created by whom, tracked/untracked', 'Drive Activity API'],
    ['File identity', 'name, Drive ID, mime type, extension, web link, icon, thumbnail', 'Drive API'],
    ['File properties', 'size, md5, sha256 (computed), created/modified time, version, revision ID, starred, shared, trashed', 'Drive API'],
    ['File capabilities', 'Full capabilities object (canEdit, canComment, canShare, etc.)', 'Drive API'],
    ['Last modifier', 'email, display name, Google user ID, photo link', 'Drive API + People API'],
    ['Owner(s)', 'email, display name per owner', 'Drive API'],
    ['Permissions', 'permission ID, type, role, email per entry', 'Drive API'],
    ['Revisions since last sync', 'revision ID, timestamp, modifier, size, md5, keep_forever', 'Drive API revisions.list'],
    ['Drive Activity', 'Action type, actor, target, timestamp, detail (rename old\u2192new, move path, etc.)', 'Drive Activity API'],
    ['Previous sync state', 'Previous revision ID, md5, size, synced timestamp', 'sync_state.json'],
    ['Status', 'UNCHANGED / MODIFIED / ADDED / DELETED', 'Computed'],
    ['Sync action', 'synced / skipped / held / rejected', 'Computed'],
    ['Validation', 'Each check: passed/failed, details', 'Computed'],
    ['Content extraction', 'Method used, success/failure, error detail', 'Pipeline'],
])

doc.add_heading('9.3 Log Security', level=2)
for s in [
    'Location: logs/activity_log.json in GitHub repo',
    'Write access: ONLY GitHub Actions bot (GITHUB_TOKEN)',
    'Branch protection: logs/ path protected via CODEOWNERS + required checks',
    'Append-only: Script only adds entries, never modifies/deletes',
    'Git history: Previous versions always recoverable, force-push blocked',
    'Read access: Admin and designated users via GitHub repo',
    'Rotation: Monthly files (activity_log_2026_02.json) to manage size',
]:
    bullet(s)

doc.add_heading('9.4 Activity Tracking', level=2)
tbl(['Event', 'Trackable?', 'API', 'Notes'], [
    ['File edited', 'YES', 'Drive API + Activity API', 'Detected at next sync'],
    ['File created', 'YES', 'Drive API + Activity API', ''],
    ['File deleted', 'YES', 'Drive API (absence)', ''],
    ['File renamed', 'YES', 'Drive Activity API', 'Old \u2192 new name captured'],
    ['File moved', 'YES', 'Drive Activity API', 'Old \u2192 new path captured'],
    ['File commented', 'YES', 'Drive Activity API', ''],
    ['Sharing changed', 'YES', 'Drive Activity API', ''],
    ['File viewed', 'NO', 'Needs Admin SDK', 'DEFERRED'],
    ['File downloaded', 'NO', 'Needs Admin SDK', 'DEFERRED'],
    ['File closed', 'NO', 'Not in any API', 'Not possible'],
])
doc.add_page_break()

# ============================================================
# 10. CONVERSION PIPELINE
# ============================================================
doc.add_heading('10. Conversion Pipeline (Smart Multi-Stage)', level=1)

doc.add_paragraph(
    'The conversion pipeline is the intelligence layer that transforms raw source files into '
    'structured KB JSON. It is derived from the proven Term 2 pipeline but refactored to: '
    '(a) run entirely on GitHub Actions, (b) skip non-Python dependencies, '
    '(c) intelligently re-run only the stages affected by each change, '
    'and (d) include automated validation as a mandatory post-build step.'
)

doc.add_heading('10.1 Change Analyzer \u2014 What Stages to Re-run', level=2)
doc.add_paragraph(
    'Before running any conversion, the Change Analyzer examines what changed since the last '
    'build and determines the minimum set of stages needed:'
)
tbl(['Change Type', 'Detection Method', 'Stages to Run', 'Automated?'], [
    ['Text on slides changed', 'md5 change + no new images in PPTX ZIP', '2, 6, 7', 'YES \u2014 fully automatic'],
    ['Speaker notes changed', 'md5 change + no new images', '2, 6, 7', 'YES'],
    ['Table modified on slide', 'md5 change + no new images', '2, 6, 7', 'YES'],
    ['New image added to slide', 'md5 change + new media files in PPTX ZIP', '1, 2, 5, 6, 7 + flag admin for Stage 4', 'PARTIAL \u2014 image description needs admin'],
    ['Image removed from slide', 'md5 change + fewer media files', '1, 5, 6, 7', 'YES'],
    ['Lesson plan docx edited', 'md5 change on .docx', '2, 6, 7', 'YES'],
    ['New lesson file added', 'New file detected', 'All stages for that lesson + flag admin', 'PARTIAL'],
    ['Native Google file edited', 'API shows new revision', '3, 6, 7', 'YES'],
    ['No changes', 'All md5 match', 'None \u2014 skip build', 'YES'],
])

doc.add_paragraph(
    'For the most common case (text, table, or speaker notes changes), the pipeline is '
    'fully automatic with zero human intervention. Only new images require admin action.'
)

doc.add_heading('10.2 Stage 1: Media Extraction', level=2)
bold_para('Script: ', 'extract_media.py (refactored \u2014 ffmpeg/video parts removed)')
bold_para('Purpose: ', 'Extract embedded images from PPTX files with slide-number tracking.')
bold_para('Method: ', 'Opens PPTX as ZIP archive, parses XML relationship files (ppt/slides/_rels/) '
          'to map images to slides, extracts media files, records slide-to-image mapping.')
bold_para('Runs when: ', 'PPTX file has new or removed images (detected by comparing media file count in ZIP).')
bold_para('Output: ', 'Extracted images organized by lesson, with metadata JSON mapping images to slides.')

doc.add_heading('10.3 Stage 2: Document Conversion', level=2)
bold_para('Script: ', 'convert_docs.py (reused from Term 2 pipeline)')
bold_para('Purpose: ', 'Convert DOCX and PPTX to structured Markdown/JSON preserving tables, headers, and notes.')
bold_para('Method: ', 'python-docx for Word files (paragraphs, tables, style-based headers). '
          'python-pptx for PowerPoint (slide-by-slide with ## Slide N headers, speaker notes, tables).')
bold_para('Runs when: ', 'Any pptx or docx file is MODIFIED or ADDED.')
bold_para('Output: ', 'Markdown files maintaining source folder hierarchy.')

doc.add_heading('10.4 Stage 3: Native Google Content Extraction', level=2)
bold_para('Script: ', 'extract_native.py (NEW \u2014 not in original pipeline)')
bold_para('Purpose: ', 'Extract structured content from Google Slides, Docs, and Sheets via their respective APIs.')
bold_para('Method: ', 'Slides API returns structured slide content (text, tables as row/col arrays, speaker notes). '
          'Docs API returns paragraphs, headings, tables. Sheets API returns cell data.')
bold_para('Runs when: ', 'A native Google file is MODIFIED.')
bold_para('Output: ', 'Structured JSON per file with full content hierarchy.')
doc.add_paragraph(
    'For files that are NOT Google-native (regular pptx/docx), this stage is skipped '
    'and Stage 2 handles extraction instead. The log notes the extraction method used.'
)

doc.add_heading('10.5 Stage 4: Image Analysis (Manual/Admin \u2014 Claude)', level=2)
bold_para('NOT automated. ', 'This stage requires Claude\u2019s multimodal capability to visually '
          'analyze images and generate educational descriptions. It cannot run in a standard GitHub Actions environment.')
doc.add_paragraph('When new images are detected:')
code("""Slack notification:
  "\U0001f5bc\ufe0f NEW IMAGES DETECTED \u2014 Admin action needed
   Lesson 4 (Term 2): 3 new images on slides 8, 12, 15
   Lesson 7 (Term 2): 1 new image on slide 3

   These images need AI descriptions before they can be
   included in the KB. Run image analysis when ready.

   Current KB build will proceed WITHOUT new image descriptions.
   Previous image data is preserved." """)
doc.add_paragraph(
    'The admin runs Claude Code agents (or Claude API) separately to describe new images. '
    'Results are added to the image descriptions batch files. On the next sync/build, '
    'the consolidation stage picks them up automatically.'
)

doc.add_heading('10.6 Stage 5: Consolidation', level=2)
bold_para('Script: ', 'consolidate.py (refactored from merge_slide_metadata.py + inline scripts)')
bold_para('Purpose: ', 'Merge all extracted content: converted markdown, image descriptions, '
          'native Google content, and metadata into a unified per-lesson structure.')
bold_para('Runs when: ', 'Any upstream stage produced new output.')

doc.add_heading('10.7 Stage 6: KB Building', level=2)
bold_para('Script: ', 'build_kb.py (refactored from compare_and_build_kb.py + fix_csv_and_add_images.py)')
bold_para('Purpose: ', 'Build the final structured JSON KB with backward-compatible schema. '
          'Original fields preserved. New enriched{} block populated per lesson.')
bold_para('Runs when: ', 'Consolidation stage produced updated content.')
bold_para('Output: ', 'term1_kb.json, term2_kb.json, term3_kb.json + changelog.md (content-level diff).')

doc.add_heading('10.8 Stage 7: Automated Validation (Post-Build Check)', level=2)
bold_para('Scripts: ', 'validation_parser.py \u2192 validation_mapper.py \u2192 validation_anomalies.py \u2192 validation_report.py')
bold_para('Purpose: ', 'Validate that every piece of content is correctly mapped to the right lesson. '
          'This is the quality gate \u2014 it runs after EVERY build.')

doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('5-Signal Consensus System:')
r.bold = True

tbl(['Signal', 'Weight', 'Method'], [
    ['Path Pattern', '1.0', 'Regex extraction from file paths ("Lesson 4" in filename, "Week 2" folder \u2192 Lessons 3-4)'],
    ['Metadata Cross-Reference', '0.95', 'Lookup against Source Inventory CSV'],
    ['Semantic Alignment', '0.8', 'Match AI-generated kb_tags against lesson keyword dictionary'],
    ['Keyword Matching', '0.7', 'Compare content themes to lesson learning objectives'],
    ['Volume Consistency', '0.5', 'Statistical validation (expected item count per lesson)'],
])

doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('Anomaly Detection:')
r.bold = True

tbl(['Anomaly Type', 'Severity', 'Trigger', 'Pipeline Action'], [
    ['MISALIGNED', 'WARNING/ERROR', 'Consensus score < 60%', 'Flag in Slack, include in report'],
    ['MISSING', 'ERROR', 'Expected content not found for a lesson', 'BLOCK KB publishing, alert admin'],
    ['DUPLICATE', 'WARNING', 'Same content mapped to unrelated lessons', 'Flag in Slack'],
    ['ORPHANED', 'WARNING', 'Content with no lesson assignment', 'Flag in Slack'],
    ['VOLUME_OUTLIER', 'INFO', 'Image count >2x standard deviation from mean', 'Log, include in report'],
    ['NAMING_INCONSISTENT', 'INFO', 'Filename doesn\u2019t match expected pattern', 'Log, include in report'],
])

doc.add_paragraph('')
bold_para('Blocking logic: ', 'If ANY anomaly with severity ERROR is detected, the KB JSON files '
          'are NOT committed to the output/ folder. The previous (good) version is preserved. '
          'Slack alert: "KB BUILD BLOCKED \u2014 validation found critical issues. See report."')

bold_para('Non-blocking: ', 'WARNING and INFO anomalies are logged and reported but do not block publishing.')

doc.add_heading('10.9 Table Preservation Strategy', level=2)
doc.add_paragraph('Tables are stored as structured arrays maintaining exact row/column alignment:')
code("""{
  "title": "Game Product Rubric (50%)",
  "source": "Lesson 4 - Rewriting the Brief.pptx",
  "slide_number": 8,
  "headers": ["Criteria", "Approaching", "Meeting", "Above Expectations"],
  "rows": [
    ["Design coherence", "Basic layout...", "Clear themed...", "Polished..."],
    ["Mechanic impl.", "Simple...", "Working...", "Complex, balanced..."]
  ],
  "row_count": 2, "col_count": 4
}

Rules: row[i][j] always corresponds to headers[j].
Merged cells: Repeated to maintain alignment.
Empty cells: "" (empty string).
Multi-line cells: Line breaks preserved.""")

doc.add_heading('10.10 Content-Level Diff', level=2)
doc.add_paragraph('After KB build, the diff engine compares new vs previous JSON:')
code("""{
  "file": "Lesson 4 - Rewriting the Brief.pptx",
  "content_diff": {
    "slides_modified": [5, 8],
    "slides_added": [13],
    "speaker_notes_changed": [5],
    "tables_modified": [{
      "slide": 8, "table_title": "Game Product Rubric",
      "changes": [
        {"type": "cell_changed", "row": 2, "col": 1,
         "old": "Basic layout", "new": "Clear themed layout"}
      ]
    }],
    "text_added": ["New activity: Team Charter exercise (slide 13)"],
    "text_modified": [{"slide": 5, "old": "teams of 3", "new": "teams of 3-4"}]
  }
}""")
doc.add_page_break()

# ============================================================
# 11. JSON SCHEMA
# ============================================================
doc.add_heading('11. JSON Schema Design (Backward-Compatible)', level=1)
doc.add_paragraph('All existing fields preserved. New data in nested enriched{} block.')

doc.add_heading('Original Fields (Preserved)', level=2)
tbl(['Field', 'Type', 'Status'], [
    ['lesson_title', 'string', 'Preserved'],
    ['url', 'string', 'Preserved'],
    ['metadata (16 sub-fields)', 'object', 'Preserved \u2014 including images array with visual descriptions'],
    ['description_of_activities', 'string', 'Preserved'],
    ['other_resources', 'string', 'Preserved'],
    ['videos_column', 'string', 'Preserved'],
    ['testing_scores', 'string', 'Preserved'],
    ['comments', 'string', 'Preserved'],
    ['prompts', 'string', 'Preserved'],
])

doc.add_heading('Enriched Block (New)', level=2)
tbl(['Field', 'Type', 'Description'], [
    ['enriched.key_facts', 'array<string>', 'Critical facts for accurate chatbot responses'],
    ['enriched.detailed_activities', 'array<object>', 'Activity breakdown: id, title, description, slide_references'],
    ['enriched.rubrics', 'array<object>', 'Full rubric tables: title, headers[], rows[][] (row/col preserved)'],
    ['enriched.teacher_notes', 'array<string>', 'Speaker notes from pptx, prefixed with slide reference'],
    ['enriched.assessment_framework', 'object', 'Weights, scoring methods, rubric types'],
    ['enriched.source_files', 'array<string>', 'Source files that contributed to this lesson'],
    ['enriched.last_updated', 'ISO 8601', 'When enriched data was last regenerated'],
    ['enriched.extraction_method', 'string', 'API or binary parse'],
    ['enriched.validation', 'object', 'Per-lesson validation result: score, anomalies, signal breakdown'],
])
doc.add_page_break()

# ============================================================
# 12. NOTIFICATIONS
# ============================================================
doc.add_heading('12. Notification System (Slack)', level=1)

doc.add_heading('12.1 Sync Notification', level=2)
code("""SYNC COMPLETE \u2014 25 Feb 2026, 00:00 UAE
  Scanned: 146 | Changed: 4 | Added: 1 | Held: 1

  SYNCED:
    Term 2 / Lesson 4.pptx \u2014 edited by ahmed@
    Term 2 / Lesson 7.pptx \u2014 edited by sarah@
  HELD:
    Term 2 / "Lesson 4.pptx" \u2014 duplicate of existing file

  Conversion pipeline starting...""")

doc.add_heading('12.2 Build + Validation Notification', level=2)
code("""KB BUILD COMPLETE \u2014 25 Feb 2026, 00:03 UAE

  Term 2 \u2014 Lesson 4:
    + 5 key facts | + 3 activities | + 1 rubric table
    ~ Slide 5 notes updated | ~ Rubric row 2 changed

  VALIDATION: PASSED (all lessons)
    12/12 lessons: content mapped correctly
    0 errors | 2 warnings | 1 info
    Warnings:
      Lesson 8: Volume outlier (42 images vs avg 26)
      Lesson 12: Naming inconsistency in source file

  Files: term2_kb.json (updated) | term1, term3 (no changes)
  Changelog: https://github.com/.../changelog.md""")

doc.add_heading('12.3 Validation BLOCKED Notification', level=2)
code("""KB BUILD BLOCKED \u2014 25 Feb 2026, 00:03 UAE

  VALIDATION FAILED: 1 critical error

  ERROR: MISSING content for Lesson 6 (Term 2)
    Expected: teacher_slides, student_slides, lesson_plan
    Found: teacher_slides only
    Possible cause: Files deleted or moved

  Previous KB version preserved. New build NOT published.
  Admin: Investigate and resolve. Re-trigger build when ready.""")

doc.add_heading('12.4 New Images Alert', level=2)
code("""NEW IMAGES DETECTED \u2014 Admin action needed

  Lesson 4 (Term 2): 3 new images (slides 8, 12, 15)
  Lesson 7 (Term 2): 1 new image (slide 3)

  KB build proceeded WITHOUT new image descriptions.
  Previous image data preserved.
  Admin: Run image analysis when ready.""")

doc.add_heading('12.5 Parent Folder Activity', level=2)
code("""PARENT FOLDER ACTIVITY
  New folder "Meeting Notes" created by fatima@
  New file "agenda.docx" added to parent by ahmed@
  These are NOT tracked. Add to config.yaml to include.""")
doc.add_page_break()

# ============================================================
# 13. GITHUB REPO
# ============================================================
doc.add_heading('13. GitHub Repository Design', level=1)
code("""curriculum-kb/  (private repo)
\u251c\u2500\u2500 sources/
\u2502   \u251c\u2500\u2500 term1/                    \u2190 synced from Google Drive
\u2502   \u251c\u2500\u2500 term2/                    \u2190 synced from Google Drive
\u2502   \u2514\u2500\u2500 term3/                    \u2190 synced from Google Drive
\u251c\u2500\u2500 output/
\u2502   \u251c\u2500\u2500 term1_kb.json             \u2190 generated
\u2502   \u251c\u2500\u2500 term2_kb.json             \u2190 generated
\u2502   \u251c\u2500\u2500 term3_kb.json             \u2190 generated
\u2502   \u251c\u2500\u2500 changelog.md              \u2190 content-level diff
\u2502   \u2514\u2500\u2500 validation_report.json    \u2190 per-build validation results
\u251c\u2500\u2500 logs/
\u2502   \u2514\u2500\u2500 activity_log_YYYY_MM.json \u2190 append-only, protected
\u251c\u2500\u2500 extracted/
\u2502   \u251c\u2500\u2500 images/                   \u2190 extracted from PPTX
\u2502   \u251c\u2500\u2500 image_descriptions/       \u2190 Claude AI descriptions
\u2502   \u2514\u2500\u2500 converted/                \u2190 markdown from conversion
\u251c\u2500\u2500 scripts/
\u2502   \u251c\u2500\u2500 sync.py                   \u2190 Drive \u2192 Git sync
\u2502   \u251c\u2500\u2500 change_analyzer.py        \u2190 determine which stages to run
\u2502   \u251c\u2500\u2500 extract_media.py          \u2190 Stage 1: image extraction
\u2502   \u251c\u2500\u2500 convert_docs.py           \u2190 Stage 2: doc conversion
\u2502   \u251c\u2500\u2500 extract_native.py         \u2190 Stage 3: Google API extraction
\u2502   \u251c\u2500\u2500 consolidate.py            \u2190 Stage 5: merge all content
\u2502   \u251c\u2500\u2500 build_kb.py               \u2190 Stage 6: JSON KB builder
\u2502   \u251c\u2500\u2500 validation_parser.py      \u2190 Stage 7a
\u2502   \u251c\u2500\u2500 validation_mapper.py      \u2190 Stage 7b
\u2502   \u251c\u2500\u2500 validation_anomalies.py   \u2190 Stage 7c
\u2502   \u251c\u2500\u2500 validation_report.py      \u2190 Stage 7d
\u2502   \u251c\u2500\u2500 diff_generator.py         \u2190 content-level diff
\u2502   \u2514\u2500\u2500 notify.py                 \u2190 Slack notifications
\u251c\u2500\u2500 config.yaml
\u251c\u2500\u2500 sync_state.json
\u251c\u2500\u2500 .github/workflows/
\u2502   \u251c\u2500\u2500 sync.yml                  \u2190 Drive sync (scheduled + manual)
\u2502   \u2514\u2500\u2500 build.yml                 \u2190 Conversion + validation
\u2514\u2500\u2500 .gitignore""")

bold_para('Two workflows: ', 'sync.yml triggers on schedule/manual. build.yml triggers when sources/ changes. '
          'Output commits don\u2019t trigger either (path filtering). No infinite loops.')
doc.add_page_break()

# ============================================================
# 14. ERROR HANDLING
# ============================================================
doc.add_heading('14. Conflict & Error Handling', level=1)
doc.add_paragraph('No Git-style conflicts. Google Drive creates sequential revisions. All logged.')

tbl(['Scenario', 'System Response', 'Recovery'], [
    ['Corrupt file', 'Rejected, others continue, alert', 'User re-uploads, next sync picks up'],
    ['Accidental deletion', 'KB keeps previous, alert', 'Restore from Drive trash (30 days) or Git history'],
    ['Duplicate filename', 'New file held, alert', 'Admin renames/deletes on Drive'],
    ['Lesson count drop', 'Term build held, alert', 'Admin investigates'],
    ['Validation ERROR', 'KB NOT published, alert', 'Admin fixes content, re-triggers build'],
    ['Validation WARNING', 'KB published, warning in report', 'Review at convenience'],
    ['Drive API failure', 'Sync fails, alert', 'Auto-retry next scheduled sync'],
    ['GitHub Actions failure', 'Build fails, email', 'Admin checks, re-triggers'],
    ['Slack webhook failure', 'Build succeeds silently', 'GitHub email as backup; fix webhook'],
    ['OAuth token expired', 'Sync fails, alert', 'Auto-refresh if refresh token valid'],
])
doc.add_page_break()

# ============================================================
# 15. TEAM WORKFLOW
# ============================================================
doc.add_heading('15. Team Workflow (Zero Extra Steps)', level=1)

doc.add_heading('Content Creators', level=2)
doc.add_paragraph('Their workflow does NOT change:')
for i, s in enumerate(['Open file on Google Drive', 'Edit in PowerPoint, Word, Google Slides, or Google Docs', 'Save', 'Done. Pipeline handles everything else.'], 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Admin', level=2)
for s in [
    'On-demand sync: GitHub Actions \u2192 Sync workflow \u2192 "Run workflow"',
    'Add term folder: Edit config.yaml tracked_folders',
    'Update Slack list: Edit config.yaml slack.notify',
    'Resolve held files: Check Slack, fix on Drive',
    'Run image analysis: Launch Claude agents for new images',
    'Review validation: Check output/validation_report.json',
    'Monitor: GitHub Actions tab for build logs',
]:
    bullet(s)

doc.add_heading('Chatbot Maintainer', level=2)
for s in [
    'Watch Slack for build notifications',
    'Pull term1/2/3_kb.json from repo',
    'Check validation_report.json for quality assurance',
    'Review changelog.md for what changed',
    'Update chatbot KB endpoint',
]:
    bullet(s)
doc.add_page_break()

# ============================================================
# 16. TESTING
# ============================================================
doc.add_heading('16. Testing & Simulation Plan', level=1)

for phase, items in [
    ('Phase 1: Pipeline Validation (Admin solo)', [
        'Set up repo, configure OAuth, verify Drive API access',
        'Run sync against real Drive folders, verify all 146 files detected and logged',
        'Verify native Google files read via APIs, Office files parsed by python-pptx/docx',
        'Run conversion pipeline, verify JSON output matches schema',
        'Run validation scripts, verify 5-signal scoring works correctly',
        'Test backward compatibility against existing Term 2 KB JSON',
        'Verify Slack notifications arrive correctly',
    ]),
    ('Phase 2: Scenario Testing (Admin + 1 colleague)', [
        'Normal edit: modify a file, verify sync + build + notification chain',
        'Duplicate detection: upload similar-named file, verify it\u2019s caught',
        'Deletion detection: delete a file, verify KB preserved + alert',
        'Corrupt file: upload broken pptx, verify graceful rejection',
        'New subfolder: create Week folder, verify detection',
        'New root folder: create folder in parent, verify notification',
        'Validation blocking: remove content, verify build is blocked',
        'Content diff: change a specific table cell, verify cell-level diff',
        'Change analyzer: modify text only, verify only Stages 2/6/7 run',
    ]),
    ('Phase 3: Guided Pilot (2-3 content creators)', [
        'Creators make real edits on Google Drive',
        'Verify Slack notifications are useful and clear',
        'Verify changes reflected in KB JSON correctly',
        'Observe where confusion arises, refine messaging',
        'Run for 1-2 weeks before full rollout',
    ]),
]:
    doc.add_heading(phase, level=2)
    for item in items:
        bullet(item)
doc.add_page_break()

# ============================================================
# 17. IMPLEMENTATION PHASES
# ============================================================
doc.add_heading('17. Implementation Phases', level=1)
tbl(['Phase', 'Deliverables', 'Dependencies', 'Priority'], [
    ['1. Google Cloud + OAuth', 'APIs enabled (DONE), OAuth working (DONE), token stored', 'Google Cloud project', 'DONE'],
    ['2. GitHub Repo Setup', 'Repo, folder structure, config.yaml, branch protection, secrets', 'Phase 1', 'Immediate'],
    ['3. Sync Pipeline', 'sync.py: detection, logging, validation, download, commit', 'Phase 2', 'High'],
    ['4. Change Analyzer', 'change_analyzer.py: determine stages to run per change type', 'Phase 3', 'High'],
    ['5. Conversion Pipeline', 'Stages 1-3, 5-6 refactored scripts: extract, convert, consolidate, build', 'Phase 3', 'High'],
    ['6. Validation Integration', 'Stage 7: parser, mapper, anomalies, report \u2014 automated post-build', 'Phase 5', 'High'],
    ['7. Diff Engine', 'diff_generator.py: content-level comparison, changelog', 'Phase 5', 'Medium'],
    ['8. Slack Notifications', 'notify.py: sync/build/validation/error/parent folder alerts', 'Phase 3', 'Medium'],
    ['9. GitHub Actions', 'sync.yml + build.yml, cron schedule, manual trigger', 'Phases 3-8', 'High'],
    ['10. Testing (3 phases)', 'Solo \u2192 pair \u2192 pilot testing', 'Phase 9', 'Critical'],
])
doc.add_page_break()

# ============================================================
# 18. APPENDIX A: CONFIG
# ============================================================
doc.add_heading('18. Appendix A: config.yaml Reference', level=1)
code("""tracked_folders:
  - name: "Term 1"
    drive_folder_id: "17s13FlHGkaNPPlf3jAUY0tSza2yxHqPe"
  - name: "Term 2"
    drive_folder_id: "16UgEwue1ROxFJyPTrowIqTQyduoNEIUb"
  - name: "Term 3"
    drive_folder_id: "1T6zzl0oqltIGcl8M4wAg2xy-z2HDZuxi"

parent_folder_id: "PARENT_FOLDER_ID"

schedule: "0 20 * * *"   # Midnight UAE

slack:
  webhook_url: "https://hooks.slack.com/services/..."
  channel: "#kb-updates"
  notify: ["@admin", "@curriculum-team"]
  admin_only: ["@admin"]

admin_emails: ["admin@company.com"]

conversion:
  extract_speaker_notes: true
  preserve_table_structure: true
  output_dir: "output/"

file_handling:
  supported_for_kb: [".pptx", ".docx"]
  supported_native: ["application/vnd.google-apps.presentation",
                     "application/vnd.google-apps.document",
                     "application/vnd.google-apps.spreadsheet"]
  store_but_skip_kb: [".pdf", ".xlsx"]
  ignore: [".mp4", ".mov", ".zip", ".wav"]
  duplicate_similarity_threshold: 0.8

logging:
  log_unchanged_files: true
  log_dir: "logs/"
  rotation: "monthly"

validation:
  block_on_error: true
  block_on_warning: false
  expected_per_lesson: ["teacher_slides"]
  lesson_keyword_dictionary: "config/lesson_keywords.yaml" """, size=7)
doc.add_page_break()

# ============================================================
# 19. APPENDIX B: LOG SCHEMA
# ============================================================
doc.add_heading('19. Appendix B: Full Log Entry Schema', level=1)
doc.add_paragraph('See v2 document Appendix B for complete JSON schema. Key addition in v3:')
code("""{
  "file_entries": [{
    ...all fields from v2...

    "content_extraction": {
      "method": "google_slides_api | binary_parse_python_pptx | none",
      "success": true,
      "error": null,
      "new_images_detected": 3,
      "images_removed": 0,
      "stages_triggered": [2, 6, 7]
    },

    "validation_result": {
      "lesson_assigned": 4,
      "confidence_score": 0.92,
      "signals": {
        "path_pattern": {"score": 1.0, "match": "Lesson 4 in filename"},
        "metadata_crossref": {"score": 0.95, "match": "CSV row 4"},
        "semantic_alignment": {"score": 0.85, "tags_matched": 5},
        "keyword_matching": {"score": 0.80, "keywords_found": 8},
        "volume_consistency": {"score": 0.50, "within_range": true}
      },
      "anomalies": []
    }
  }]
}""")
doc.add_page_break()

# ============================================================
# 20. APPENDIX C: VALIDATION SIGNALS
# ============================================================
doc.add_heading('20. Appendix C: Validation Signal Definitions', level=1)
doc.add_paragraph('Lesson keyword dictionary used by Signal 3 (Semantic) and Signal 4 (Keyword):')
tbl(['Lesson', 'Keywords'], [
    ['1', 'design brief, problem statement, audience, UAE heritage, cultural context, sustainability'],
    ['2', 'persona, empathy map, UX, player needs, motivations, user research'],
    ['3', 'primary research, secondary research, AI research, bias, sources, data collection'],
    ['4', 'design specification, team roles, constraints, success criteria, team charter'],
    ['5', 'brainstorming, concept generation, micro-prototype, storyboard, ideation'],
    ['6', 'prototype, core mechanic, debugging, testing, iteration, gameplay loop'],
    ['7', 'gameplay expansion, immersion, visuals, sound, dialogue, polish'],
    ['8', 'peer testing, WWW/EBI, feedback analysis, theme mapping, playtesting'],
    ['9', 'iteration, refinement, feedback implementation, impact vs effort matrix'],
    ['10', 'team roles, project manager, milestones, timeline, risk management'],
    ['11', 'documentation, portfolio, evidence, curation, reflection, showcase'],
    ['12', 'reflection, evaluation, SMART goals, Term 3 preview, progress review'],
])

doc.add_paragraph('')
doc.add_paragraph('Note: This dictionary is for Term 2. Term 1 and Term 3 will have their own dictionaries '
                   'configured in config/lesson_keywords.yaml as those terms are integrated.')

# ============================================================
# END
# ============================================================
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\u2014\u2014\u2014 End of Document \u2014\u2014\u2014')
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99); r.font.size = Pt(9)

doc.save('KB_Maintenance_Mechanism_Pipeline_v3.docx')
print('Document saved: KB_Maintenance_Mechanism_Pipeline_v3.docx')
