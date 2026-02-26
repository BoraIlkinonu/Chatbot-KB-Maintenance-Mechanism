"""
Generate comprehensive pipeline documentation as DOCX.
Includes all discussion points, decisions, folder analysis, and technical details.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def add_code_block(doc, lines, font_size=8):
    for line in lines if isinstance(lines, list) else lines.split('\n'):
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(font_size)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    doc.add_paragraph('')

def add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.rows[i+1].cells[j].text = str(cell)
            for p in table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph('')

def add_bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(5):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Curriculum KB Maintenance Mechanism')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Comprehensive Pipeline Design & Technical Specification')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('25 February 2026 | Version 2.0')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Classification: Internal | Based on Comprehensive Design Discussion')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc = [
    '1. Executive Summary',
    '2. Problem Statement & QA Impact',
    '3. Design Principles & Key Decisions',
    '4. Source Folder Analysis (Live Scan Results)',
    '5. Complete Pipeline Architecture',
    '6. Google Drive Integration Layer',
    '   6.1 APIs & Permissions',
    '   6.2 OAuth Authentication',
    '   6.3 Change Detection',
    '   6.4 File Format Handling (Native vs Office)',
    '7. Sync Layer — Detailed Design',
    '   7.1 Trigger Modes',
    '   7.2 Step-by-Step Sync Process',
    '   7.3 Validation Checks',
    '   7.4 File Scenarios (Replace, Duplicate, Delete, New Folder)',
    '8. Comprehensive Logging System',
    '   8.1 Design Philosophy: Log Everything',
    '   8.2 Log Entry Schema (Exhaustive)',
    '   8.3 Log Security (Read-Only Protection)',
    '   8.4 Activity Tracking Capabilities & Limitations',
    '9. GitHub Repository Design',
    '   9.1 Repository Structure',
    '   9.2 Branch Protection',
    '   9.3 GitHub Actions Workflow',
    '10. Conversion Pipeline',
    '   10.1 Extraction Strategy (pptx/docx/native)',
    '   10.2 Table Preservation',
    '   10.3 Content-Level Diff',
    '11. JSON Schema Design (Backward-Compatible)',
    '   11.1 Original Fields (Preserved)',
    '   11.2 Enriched Block (New)',
    '   11.3 Rubric Table Structure',
    '   11.4 Sample Output',
    '12. Notification System (Slack)',
    '   12.1 Activity Alerts',
    '   12.2 Build Notifications',
    '   12.3 Error & Warning Alerts',
    '   12.4 Parent Folder Monitoring',
    '13. Conflict & Error Handling',
    '14. Team Workflow (Zero Extra Steps)',
    '15. Testing & Simulation Plan',
    '16. Implementation Phases',
    '17. Appendix A: config.yaml Reference',
    '18. Appendix B: Full Log Entry Schema',
    '19. Appendix C: Google Cloud Project Setup',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)

doc.add_page_break()

# ============================================================
# SECTION 1: EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'This document defines the complete design for an automated Curriculum Knowledge Base (KB) '
    'maintenance mechanism. The system monitors Google Drive folders containing curriculum source '
    'files (pptx, docx, Google Slides, Google Docs), detects changes, validates them, converts '
    'content to structured JSON, and notifies the team via Slack \u2014 all without requiring any '
    'change to how the curriculum team currently works.'
)

doc.add_paragraph(
    'The mechanism was designed in response to QA findings (11 February 2026) that identified '
    'stale and incomplete KB data as the root cause of false-positive hallucination flags in the '
    'Endstar AI Assistant chatbot. The chatbot serves students in the Explorer\u2019s Programme '
    '(G9\u2013G10, UAE) and draws knowledge from two separate pipelines: our curriculum team\u2019s '
    'KB and the Endstar platform team\u2019s KB. This mechanism addresses our pipeline exclusively.'
)

p = doc.add_paragraph()
r = p.add_run('Core Design Principle: ')
r.bold = True
p.add_run(
    'Users edit files on Google Drive exactly as they do today. No Git, no GitHub Desktop, '
    'no extra steps, no locking, no restrictions. The entire pipeline is invisible to content '
    'creators. All intelligence is in the automated sync and validation layer.'
)

doc.add_page_break()

# ============================================================
# SECTION 2: PROBLEM STATEMENT
# ============================================================
doc.add_heading('2. Problem Statement & QA Impact', level=1)

doc.add_heading('Current Problems', level=2)
problems = [
    'Stale KB data: The curriculum KB was generated once from a CSV summary and never updated, even as source lesson plans evolved on Google Drive.',
    'Mixed content: Curriculum data was mixed with Endstar technical data that belongs to the other team\u2019s pipeline.',
    'Shallow extraction: The CSV-to-JSON conversion captured metadata-level info but missed deep content \u2014 rubric table descriptors, activity breakdowns, teacher notes from speaker notes, and assessment framework structures.',
    'No update mechanism: Changes to lesson plans on Google Drive have no path to reach the chatbot KB. The pipeline was run once and never again.',
    'No audit trail: No record of who changed what, when, or why. No visibility into KB data freshness.',
]
for p_text in problems:
    add_bullet(doc, p_text)

doc.add_heading('QA Impact (11 February 2026 Report)', level=2)
add_table(doc,
    ['Category', 'Score', 'Root Cause'],
    [
        ['Term 1 Knowledge', '99.1%', 'KB adequate for Term 1'],
        ['Term 2 Knowledge', '75.7%', 'Missing rubric descriptors, incomplete assessment data'],
        ['Cross-Term Confusion', '69.6%', 'KB lacks structural awareness of term differences'],
        ['Student Simulation', '75.6%', 'Fabricated platform details due to missing KB content'],
        ['Hallucination Probing', '98.0%', 'Chatbot correctly defers when unsure'],
        ['Overall Factual Accuracy', '72.7%', 'KB gaps, not chatbot defects'],
        ['Overall Hallucination Resistance', '71.8%', 'False positives from incomplete KB ground truth'],
    ]
)

doc.add_paragraph(
    'The QA report explicitly recommended a \u201cKB Synchronisation Mechanism\u201d as Priority 1. '
    'A significant portion of low scores were caused by KB gaps, not actual chatbot defects. '
    'Fixing the KB pipeline is expected to improve multiple scores significantly.'
)

doc.add_page_break()

# ============================================================
# SECTION 3: DESIGN PRINCIPLES
# ============================================================
doc.add_heading('3. Design Principles & Key Decisions', level=1)

decisions = [
    ('Google Drive is the source of truth', 'Users work on Google Drive. They do not interact with Git, GitHub, or any new tool. The pipeline comes to them, not the other way around.'),
    ('No file locking or ownership restrictions', 'Anyone can edit any file at any time. Git-style file locking is not bulletproof for binary files, and adding a Google Sheet lock register was rejected as impractical (extra step nobody would follow). Instead, the system detects and handles issues after the fact.'),
    ('Log everything, restrict nothing', 'Every file, every sync cycle, every user action is logged exhaustively. UNCHANGED files are logged as unchanged. The log is the audit trail and the backbone of the system. It is append-only and protected (read-only for everyone except the automated pipeline).'),
    ('Backward-compatible JSON schema', 'The existing KB JSON structure is preserved exactly. New enriched data is added in a nested "enriched" block per lesson. Systems reading old fields continue to work without changes.'),
    ('Scheduled + on-demand triggers', 'Default: midnight UAE daily sync. Admin can trigger on-demand for urgent updates. No real-time automatic triggers to avoid excessive builds.'),
    ('Slack for all notifications', 'Activity alerts, build notifications, error warnings, and parent folder monitoring all go to a Slack channel. Distribution list is configurable.'),
    ('No user workflow changes', 'The entire mechanism is invisible to content creators. Their workflow remains: open file on Google Drive, edit, save. Nothing else.'),
    ('Comprehensive validation without blocking', 'The pipeline validates every change (duplicates, deletions, integrity, lesson count) but only HOLDS suspect changes for admin review \u2014 it never silently drops or blocks legitimate edits.'),
    ('Content-level diff after conversion', 'Google Drive API provides metadata-level changes (who, when, size). Content-level diff (which slide changed, which table cell was modified) is computed after conversion by comparing previous and current extracted JSON.'),
]
for title, desc in decisions:
    p = doc.add_paragraph()
    r = p.add_run(f'{title}: ')
    r.bold = True
    p.add_run(desc)

doc.add_page_break()

# ============================================================
# SECTION 4: SOURCE FOLDER ANALYSIS
# ============================================================
doc.add_heading('4. Source Folder Analysis (Live Scan Results)', level=1)

doc.add_paragraph(
    'The following analysis is based on a live scan of the three Google Drive folders performed '
    'on 25 February 2026 using the Google Drive API with OAuth authentication.'
)

doc.add_heading('Overview', level=2)
add_table(doc,
    ['Folder', 'Content', 'Files', 'Subfolders', 'Size'],
    [
        ['Term 1 Content (Foundations)', 'January 2026 Cohort', '51', '14', '481 MB'],
        ['Term 2 Content (Accelerator)', 'August 2025 Cohort', '75', '45', '1,069 MB'],
        ['Term 3 Content (Mastery)', 'August 2025 Cohort', '20', '15', '510 MB'],
        ['TOTAL', '', '146', '74', '2,060 MB'],
    ]
)

doc.add_heading('Term 1 Structure', level=2)
add_code_block(doc, """Term 1 Content (Foundations) - January 2026 Cohort/
\u251c\u2500\u2500 Administrative/
\u2502   \u251c\u2500\u2500 Explorers Programme Curriculum Specifications Term 1 (Google Doc)
\u2502   \u2514\u2500\u2500 IB Wrapper Design Curriculum Mapping.pptx
\u251c\u2500\u2500 Feedback & Surveys/
\u2502   \u2514\u2500\u2500 Feedback Forms.docx
\u251c\u2500\u2500 Lessons Resources/
\u2502   \u251c\u2500\u2500 Lesson Decks (Slides)/
\u2502   \u2502   \u251c\u2500\u2500 Explorer's Programme - Lesson 1.pptx  (16.2MB)
\u2502   \u2502   \u251c\u2500\u2500 Explorer's Programme - Lesson 2.pptx  (11.5MB)
\u2502   \u2502   \u251c\u2500\u2500 ... (Lessons 3-22, all .pptx)
\u2502   \u2502   \u2514\u2500\u2500 Lesson Decks (Slides) (Google Slides shortcut)
\u2502   \u2514\u2500\u2500 Student Assessment Templates/
\u2502       \u251c\u2500\u2500 Pitch Rubric.pptx (Google Slides native)
\u2502       \u251c\u2500\u2500 Students Assessment Templates (Google Doc)
\u2502       \u251c\u2500\u2500 Students Master Checklist (Google Sheet)
\u2502       \u2514\u2500\u2500 US Curriculum Assessment Guide.pdf
\u251c\u2500\u2500 Programme Overview/
\u2502   \u251c\u2500\u2500 Explorer's Programme Overview.pdf (34.2MB)
\u2502   \u251c\u2500\u2500 Learning Outcomes - Jan 2026 Cohort (Google Doc)
\u2502   \u2514\u2500\u2500 Teacher Training - Explorer Programme.pdf (36.2MB)
\u251c\u2500\u2500 Quick Access Links/
\u251c\u2500\u2500 Showcase & Deliverables/
\u251c\u2500\u2500 Teacher Resources/
\u2502   \u251c\u2500\u2500 Exemplar Workbook/ (2 PDFs, 26MB each)
\u2502   \u251c\u2500\u2500 Endstar Level Design - Teacher Rubric (Google Doc)
\u2502   \u2514\u2500\u2500 Learning Schedule.xlsx
\u2514\u2500\u2500 Technical Support/ (Setup guides, troubleshooting, 98MB video)""", font_size=7)

doc.add_heading('Key Observations: Term 1', level=3)
observations_t1 = [
    '22 lesson decks (Lessons 1-22) as .pptx files, 7.5\u201318.7MB each',
    'No separate Teacher Slides vs Student Slides \u2014 single set of lesson decks',
    'No lesson plan documents (docx) \u2014 only slide decks',
    'Mix of native Google formats (Docs, Sheets, Slides) and Office formats (pptx, docx, xlsx)',
    'Assessment templates include a Google Sheet (Students Master Checklist) that the Sheets API can read directly',
    'Large video file (98.8MB .mov) in Technical Support \u2014 not relevant for KB extraction',
]
for o in observations_t1:
    add_bullet(doc, o)

doc.add_heading('Term 2 Structure', level=2)
add_code_block(doc, """Term 2 Content (Accelerator) - August 2025 Cohort/
\u251c\u2500\u2500 Teacher Resources/
\u2502   \u251c\u2500\u2500 Assessment Guides/
\u2502   \u2502   \u251c\u2500\u2500 Exemplar work/ (Weeks 1-6, pptx + mp4 files)
\u2502   \u2502   \u251c\u2500\u2500 Student Guide/ (docx + pdf)
\u2502   \u2502   \u2514\u2500\u2500 Teacher Guide/ (docx + pdf)
\u2502   \u251c\u2500\u2500 Curriculum Alignment/ (PDF)
\u2502   \u251c\u2500\u2500 Curriculum Content/
\u2502   \u2502   \u251c\u2500\u2500 LO's and Success Criteria/ (PDF)
\u2502   \u2502   \u251c\u2500\u2500 Week 1/
\u2502   \u2502   \u2502   \u251c\u2500\u2500 Lesson Plans/     (Lesson 1.docx, Lesson 2.docx)
\u2502   \u2502   \u2502   \u251c\u2500\u2500 Students Slides/  (Google Slides native)
\u2502   \u2502   \u2502   \u2514\u2500\u2500 Teachers Slides/  (Lesson 1.pptx, Lesson 2.pptx)
\u2502   \u2502   \u251c\u2500\u2500 Week 2/ ... Week 6/ (same structure)
\u2502   \u2502   \u2514\u2500\u2500 NOTE: Student Slides = Google native, Teacher Slides = pptx
\u2502   \u251c\u2500\u2500 Design Briefs and Exemplar Games/ (6 PDFs)
\u2502   \u251c\u2500\u2500 Professional Development/ (3 PDFs)
\u2502   \u251c\u2500\u2500 Student Portfolio/ (1 large Google Slides, 34.6MB)
\u2502   \u251c\u2500\u2500 Curriculum Specifications Term 2.docx
\u2502   \u251c\u2500\u2500 Learning Schedule (Google Sheet)
\u2502   \u2514\u2500\u2500 Term 2 Introduction content.pptx
\u2514\u2500\u2500 Technical Support Guides/ (Classkick, Endstar, School AI)""", font_size=7)

doc.add_heading('Key Observations: Term 2', level=3)
observations_t2 = [
    '12 lessons (Weeks 1-6, 2 lessons per week) \u2014 richest structure of all three terms',
    'THREE types of content per lesson: Lesson Plans (docx), Teacher Slides (pptx), Student Slides (Google Slides native)',
    'Teacher Slides and Student Slides appear to be duplicates (same file size) \u2014 pipeline should detect and handle this',
    'Lesson Plan documents (docx) are separate from slide decks \u2014 valuable for structured text extraction',
    'Exemplar Work exists for each week pair (Lessons 1-2, 3-4, etc.)',
    'Assessment guides exist as both docx and PDF (visual version) \u2014 Student Guide and Teacher Guide',
    'Naming inconsistency: "Lesson Plans" (plural) vs "Lesson Plan" (singular) in different weeks',
    'Naming inconsistency: "Exampler" (typo) instead of "Exemplar"',
    'Some files are Google-native (Student Slides, Learning Schedule, Activities & Portfolio Deck)',
]
for o in observations_t2:
    add_bullet(doc, o)

doc.add_heading('Term 3 Structure', level=2)
add_code_block(doc, """Term 3 Content (Mastery) - August 2025 Cohort/
\u251c\u2500\u2500 Teacher Resources/
\u2502   \u251c\u2500\u2500 Assessment/ (empty)
\u2502   \u251c\u2500\u2500 Curriculum Alignment/ (empty)
\u2502   \u251c\u2500\u2500 Curriculum Content/
\u2502   \u2502   \u251c\u2500\u2500 Lesson supporting videos/ (4 videos, 230MB total)
\u2502   \u2502   \u251c\u2500\u2500 Week 1/Teacher Slides/ (Lesson 1.pptx, Lesson 2.pptx)
\u2502   \u2502   \u251c\u2500\u2500 Week 2/ (Lesson 3.pptx, Lesson 4.pptx)
\u2502   \u2502   \u251c\u2500\u2500 Week 3/ (Lesson 5.pptx, Lesson 6.pptx)
\u2502   \u2502   \u251c\u2500\u2500 Week 4/ (Lesson 7.pptx, Lesson 8.pptx)
\u2502   \u2502   \u251c\u2500\u2500 Week 5/ (Lesson 9.pptx, Lesson 10.pptx + 2 videos)
\u2502   \u2502   \u251c\u2500\u2500 Week 6/ (Lesson 11.pptx, Lesson 12.pptx)
\u2502   \u2502   \u251c\u2500\u2500 Week 7/ (Lesson 13.pptx, Lesson 14.pptx)
\u2502   \u2502   \u2514\u2500\u2500 Week 8/ (empty)
\u2514\u2500\u2500 Technical Support Guides/ (empty)""", font_size=7)

doc.add_heading('Key Observations: Term 3', level=3)
observations_t3 = [
    '14 lessons (Weeks 1-7, 2 per week), Week 8 empty \u2014 possibly in development',
    'Only Teacher Slides (pptx), no separate Student Slides or Lesson Plan documents',
    'No assessment guides or exemplar work yet',
    'Assessment and Curriculum Alignment folders exist but are empty',
    'Videos mixed into curriculum content folders (not separated like Term 2)',
    'Simpler structure than Term 2 \u2014 may still be under development',
]
for o in observations_t3:
    add_bullet(doc, o)

doc.add_heading('File Type Distribution Across All Terms', level=2)
add_table(doc,
    ['File Type', 'Count', 'Total Size', 'KB Extraction Method'],
    [
        ['PowerPoint (.pptx)', '~50', '~800MB', 'python-pptx (binary parse) or Slides API if native'],
        ['Google Slides (native)', '~8', 'N/A (cloud)', 'Google Slides API (structured JSON)'],
        ['Word (.docx)', '~15', '~5MB', 'python-docx (binary parse) or Docs API if native'],
        ['Google Docs (native)', '~8', 'N/A (cloud)', 'Google Docs API (structured JSON)'],
        ['Google Sheets (native)', '~3', 'N/A (cloud)', 'Google Sheets API (cell data)'],
        ['PDF', '~25', '~170MB', 'Stored, not parsed for KB (reference material)'],
        ['Video (.mp4, .mov)', '~10', '~750MB', 'NOT extracted \u2014 logged only, ignored for KB'],
        ['Excel (.xlsx)', '1', '46KB', 'openpyxl or Sheets API if converted'],
    ]
)

doc.add_heading('Naming Inconsistencies Detected', level=2)
add_table(doc,
    ['Issue', 'Example', 'Impact on Pipeline'],
    [
        ['Trailing spaces in filenames', '"Lesson 1 .pptx" (space before dot)', 'Fuzzy match must normalize whitespace'],
        ['Typos', '"Exampler Work" should be "Exemplar Work"', 'Human error \u2014 log and flag but accept'],
        ['Singular vs plural folders', '"Lesson Plans" vs "Lesson Plan"', 'Folder scanning must be flexible'],
        ['Inconsistent separators', '"Lesson 1" vs "Lesson1" vs "Lesson_1"', 'Normalization in duplicate check'],
        ['Missing lesson titles in name', 'Term 3: just "Lesson 5.pptx" vs Term 1: "Explorer\'s Programme - Lesson 5.pptx"', 'Lesson number extraction must handle both formats'],
        ['Mixed case', 'Some files capitalized differently', 'Case-insensitive comparison'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 5: COMPLETE PIPELINE ARCHITECTURE
# ============================================================
doc.add_heading('5. Complete Pipeline Architecture', level=1)

add_code_block(doc, """\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502  GOOGLE DRIVE (Source of Truth)                              \u2502
\u2502                                                              \u2502
\u2502  \U0001f4c1 Term 1 Content (Foundations) - 22 lessons             \u2502
\u2502  \U0001f4c1 Term 2 Content (Accelerator) - 12 lessons             \u2502
\u2502  \U0001f4c1 Term 3 Content (Mastery) - 14 lessons                 \u2502
\u2502                                                              \u2502
\u2502  Anyone can edit. No restrictions. No extra steps.           \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
               \u2502
          Google Drive API + Slides/Docs/Sheets APIs
          (OAuth, refresh token, 6 API scopes)
               \u2502
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502  SYNC + LOGGING LAYER (GitHub Actions cloud runner)          \u2502
\u2502                                                              \u2502
\u2502  Step 1: DETECT changes (Drive API vs sync_state.json)       \u2502
\u2502  Step 2: LOG everything (append-only, read-only protected)   \u2502
\u2502  Step 3: VALIDATE (duplicate, deletion, integrity, count)    \u2502
\u2502  Step 4: DOWNLOAD accepted files (to cloud runner, NOT local)\u2502
\u2502  Step 5: COMMIT to Git repo                                  \u2502
\u2502  Step 6: NOTIFY via Slack                                    \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
               \u2502
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502  CONVERSION PIPELINE                                        \u2502
\u2502                                                              \u2502
\u2502  Native Google files \u2192 Slides/Docs/Sheets API \u2192 JSON        \u2502
\u2502  Office files (pptx/docx) \u2192 python-pptx/docx \u2192 JSON        \u2502
\u2502  Previous JSON vs New JSON \u2192 Content-level diff             \u2502
\u2502  Output: term1_kb.json, term2_kb.json, term3_kb.json         \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518
               \u2502
\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510
\u2502  OUTPUTS & NOTIFICATIONS                                    \u2502
\u2502                                                              \u2502
\u2502  Slack: Activity alerts, build notifications, errors         \u2502
\u2502  JSON: term1_kb.json, term2_kb.json, term3_kb.json           \u2502
\u2502  Changelog: What changed inside each file (content diff)     \u2502
\u2502  Logs: Complete audit trail of every action                  \u2502
\u2502  Delivery: Chatbot maintainer pulls JSONs from repo          \u2502
\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518""")

doc.add_page_break()

# ============================================================
# SECTION 6: GOOGLE DRIVE INTEGRATION
# ============================================================
doc.add_heading('6. Google Drive Integration Layer', level=1)

doc.add_heading('6.1 APIs & Permissions', level=2)
add_table(doc,
    ['#', 'API', 'Scope', 'Purpose', 'Admin Required?'],
    [
        ['1', 'Drive API v3', 'drive.readonly', 'File listing, download, metadata, revisions, permissions, checksums, folder structure', 'No'],
        ['2', 'Drive Activity API v2', 'drive.activity.readonly', 'Detailed action feed: created/edited/moved/renamed/deleted/commented/shared', 'No'],
        ['3', 'Slides API v1', 'presentations.readonly', 'Structured content from Google Slides: slides, text, tables, speaker notes', 'No'],
        ['4', 'Docs API v1', 'documents.readonly', 'Structured content from Google Docs: paragraphs, tables, headings, lists', 'No'],
        ['5', 'Sheets API v4', 'spreadsheets.readonly', 'Cell data from Google Sheets: values, ranges, formatting', 'No'],
        ['6', 'People API', 'directory.readonly', 'Resolve user IDs to full profiles: name, email, photo, department', 'No'],
        ['7', 'Admin SDK Reports (DEFERRED)', 'admin.reports.audit.readonly', 'View/download events \u2014 requires Workspace admin, to be added later', 'YES'],
    ]
)

doc.add_heading('6.2 OAuth Authentication', level=2)
doc.add_paragraph(
    'Authentication uses OAuth 2.0 with an installed application flow. The user authorizes once '
    'via browser; a refresh token is stored securely as a GitHub Actions secret for subsequent runs.'
)
auth_steps = [
    'Google Cloud Project: poetic-dock-483707-n1',
    'Client type: Installed application (Desktop)',
    'First run: Opens browser for Google consent screen',
    'Token stored: As GitHub Actions encrypted secret (GOOGLE_TOKEN_JSON)',
    'Refresh: Automatic using refresh token \u2014 no re-login needed',
    'Scopes: All 6 non-admin scopes requested at once during initial auth',
]
for s in auth_steps:
    add_bullet(doc, s)

doc.add_heading('6.3 Change Detection', level=2)
doc.add_paragraph(
    'The sync script maintains a sync_state.json file in the repo that records the last-known '
    'state of every file. On each sync run, it compares the current Drive state against this file '
    'to detect changes.'
)
add_code_block(doc, """{
  "last_sync": "2026-02-17T20:00:00Z",
  "files": {
    "1aBcDeFg": {
      "name": "Lesson 4 - Rewriting the Brief.pptx",
      "folder": "Term 2",
      "revision_id": "r1232",
      "md5_checksum": "abc123...",
      "size_bytes": 4518000,
      "modified_time": "2026-02-16T14:00:00Z"
    }
  }
}

Comparison logic:
  File in Drive + File in sync_state + same md5     \u2192 UNCHANGED
  File in Drive + File in sync_state + different md5 \u2192 MODIFIED
  File in Drive + NOT in sync_state                  \u2192 ADDED
  NOT in Drive + File in sync_state                  \u2192 DELETED""")

doc.add_heading('6.4 File Format Handling', level=2)
add_code_block(doc, """For each file detected:

\u251c\u2500\u2500 Google Slides (native)
\u2502   \u251c\u2500\u2500 Use Slides API \u2192 structured JSON (slides, text, tables, notes)
\u2502   \u251c\u2500\u2500 Also export as .pptx for Git storage (backup)
\u2502   \u2514\u2500\u2500 Log: "content_extraction_method": "google_slides_api"
\u2502
\u251c\u2500\u2500 Google Docs (native)
\u2502   \u251c\u2500\u2500 Use Docs API \u2192 structured JSON (paragraphs, tables, headings)
\u2502   \u251c\u2500\u2500 Also export as .docx for Git storage (backup)
\u2502   \u2514\u2500\u2500 Log: "content_extraction_method": "google_docs_api"
\u2502
\u251c\u2500\u2500 Google Sheets (native)
\u2502   \u251c\u2500\u2500 Use Sheets API \u2192 cell data, ranges, formatting
\u2502   \u251c\u2500\u2500 Also export as .xlsx for Git storage (backup)
\u2502   \u2514\u2500\u2500 Log: "content_extraction_method": "google_sheets_api"
\u2502
\u251c\u2500\u2500 PowerPoint (.pptx)
\u2502   \u251c\u2500\u2500 Download binary from Drive API
\u2502   \u251c\u2500\u2500 Parse with python-pptx on GitHub Actions runner
\u2502   \u2514\u2500\u2500 Log: "content_extraction_method": "binary_parse_python_pptx"
\u2502
\u251c\u2500\u2500 Word (.docx)
\u2502   \u251c\u2500\u2500 Download binary from Drive API
\u2502   \u251c\u2500\u2500 Parse with python-docx on GitHub Actions runner
\u2502   \u2514\u2500\u2500 Log: "content_extraction_method": "binary_parse_python_docx"
\u2502
\u251c\u2500\u2500 PDF / Video / Other
\u2502   \u251c\u2500\u2500 Download and store in Git (for reference)
\u2502   \u251c\u2500\u2500 NOT parsed for KB extraction
\u2502   \u2514\u2500\u2500 Log: "content_extraction_method": "none_unsupported_format"
""")

doc.add_page_break()

# ============================================================
# SECTION 7: SYNC LAYER
# ============================================================
doc.add_heading('7. Sync Layer \u2014 Detailed Design', level=1)

doc.add_heading('7.1 Trigger Modes', level=2)
add_table(doc,
    ['Mode', 'Trigger', 'Who Can Trigger', 'When to Use'],
    [
        ['Scheduled', 'Cron: midnight UAE (20:00 UTC)', 'Automatic', 'Default daily sync \u2014 picks up all changes from the day'],
        ['On-demand', 'GitHub Actions "Run workflow" button', 'Admin only', 'Urgent updates, testing, ad-hoc sync'],
    ]
)

doc.add_heading('7.2 Step-by-Step Sync Process', level=2)

steps = [
    ('Step 1: DETECT changes',
     'For each tracked folder in config.yaml, call Drive API to list all files and their current state '
     '(revision ID, md5, size, modified time, last modifying user). Compare against sync_state.json. '
     'Categorize every file as UNCHANGED, MODIFIED, ADDED, or DELETED. Also scan parent folder for '
     'any new or untracked folders/files and report them.'),
    ('Step 2: LOG everything',
     'For every file (including UNCHANGED), capture the complete API response: user details, file metadata, '
     'revision history, permissions, capabilities, checksums, links, sizes, timestamps \u2014 everything the '
     'API returns, no filtering. Use Drive Activity API to get detailed action history (edits, renames, '
     'moves, comments, shares). Use People API to resolve user profiles. Append all entries to '
     'logs/activity_log.json. This file is append-only and write-protected (only GitHub Actions bot can modify).'),
    ('Step 3: VALIDATE',
     'Run validation checks on MODIFIED, ADDED, and DELETED files only (UNCHANGED files skip validation). '
     'Each check either ACCEPTS or HOLDS the file. Held files are not synced but are logged and reported. '
     'See Section 7.3 for detailed validation checks.'),
    ('Step 4: DOWNLOAD accepted files',
     'For each ACCEPTED file: download from Drive API to the GitHub Actions cloud runner temporary filesystem. '
     'For native Google files: also extract structured content via Slides/Docs/Sheets APIs. '
     'Place files in sources/termN/ directory matching the Drive folder structure. '
     'Note: "local" means the cloud runner, NOT anyone\'s personal computer. No local machines are involved.'),
    ('Step 5: COMMIT to Git',
     'Stage all changed files in sources/, logs/activity_log.json, and sync_state.json. '
     'Create a commit with a descriptive message listing all synced files and who changed them. '
     'Push to main. If nothing changed (all files UNCHANGED), skip commit and do not send notification.'),
    ('Step 6: NOTIFY via Slack',
     'Send Slack message to configured channel with: list of all synced files, who changed each file, '
     'any validation warnings (held files, duplicates detected, new folders), and a link to the repo. '
     'If the conversion pipeline follows, a second notification is sent when the KB build completes.'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    r = p.add_run(f'{title}: ')
    r.bold = True
    p.add_run(desc)

doc.add_heading('7.3 Validation Checks', level=2)

checks = [
    ('DUPLICATE CHECK (3 layers)',
     'Layer 1 \u2014 Exact name match: Detect "(1)", "(2)" copies created by Google Drive\'s "Keep both" option. '
     'Layer 2 \u2014 Fuzzy name match: Normalize filenames (lowercase, remove underscores/hyphens/extra spaces), '
     'extract lesson numbers, compare with Levenshtein distance. Flag if similarity > 0.8 or same lesson number. '
     'Layer 3 \u2014 Content comparison: If fuzzy match triggers, download both files and compare MD5 checksums. '
     'If identical content: "These files are exact duplicates." If different: "Two different files claim to be Lesson N."'),
    ('DELETION CHECK',
     'If a file existed in sync_state.json but is no longer in Drive: DO NOT remove from Git sources/. '
     'Alert admin: "File X was deleted from Drive by user@. KB NOT updated \u2014 admin must confirm removal." '
     'Google Drive keeps deleted files in Trash for 30 days, so recovery is possible.'),
    ('INTEGRITY CHECK',
     'Download the file and attempt to parse with python-pptx (for .pptx) or python-docx (for .docx). '
     'For native Google files, attempt API access. If parsing fails: REJECT this file, do NOT include in '
     'sync, alert: "File X could not be parsed \u2014 may be corrupt." Continue syncing other files.'),
    ('LESSON COUNT CHECK',
     'After all files are processed, count total parseable lesson files per term. Compare against '
     'previous sync count. If the new count is LOWER: HOLD the entire term build and alert admin. '
     'This catches accidental bulk deletions or folder restructuring that removes content.'),
]
for title, desc in checks:
    p = doc.add_paragraph()
    r = p.add_run(f'{title}: ')
    r.bold = True
    p.add_run(desc)

doc.add_heading('7.4 File Scenarios', level=2)

scenarios = [
    ('Normal edit (most common)', 'User edits and saves a file on Drive. Drive creates a new revision (same file ID). Sync detects MODIFIED via md5 change. Downloads latest revision. Commits. Logs who, when, what revision.'),
    ('File replaced via web upload ("Replace")', 'User uploads a file with the same name and chooses "Replace." Drive creates a new revision of the existing file. Identical to normal edit from the pipeline\'s perspective.'),
    ('File uploaded with "Keep both"', 'Google Drive creates "Filename (1).pptx" as a separate file. Duplicate check Layer 1 detects the "(1)" pattern. File is HELD. Slack alert sent with both filenames.'),
    ('File with similar name (typo)', 'New file "Lesson 4.pptx" uploaded when "Lesson 4 - Rewriting the Brief.pptx" exists. Duplicate check Layer 2 detects same lesson number. File is HELD pending admin review.'),
    ('File deleted', 'File removed from Drive. Deletion check catches it. KB keeps previous version. Admin must confirm removal.'),
    ('New file added (new lesson)', 'A new pptx appears that doesn\'t match any existing file. If naming convention is valid: ACCEPTED and synced. Slack notification: "New file detected: Lesson 13.pptx added to Term 2."'),
    ('New tracked-level folder added (e.g., Term 4)', 'New folder created at the parent level. Not in config.yaml tracked_folders. Slack notification: "New folder \'Term 4 Content\' detected. NOT tracked. Admin: add to config.yaml to include."'),
    ('New subfolder inside tracked term', 'New subfolder (e.g., "Week 7/") created inside a tracked term. Automatically scanned for files. Slack notification: "New subfolder \'Week 7\' added to Term 3."'),
    ('Unrelated folder/file in parent', 'File or folder added to the parent directory that is not a term folder. Slack notification: "Activity in parent folder: \'Meeting Notes\' folder created by user@." Logged but ignored for KB.'),
    ('Corrupt file pushed', 'File that python-pptx/python-docx cannot parse. Integrity check catches it. File is REJECTED. Other files continue syncing. Slack alert: "Lesson 5.pptx could not be parsed."'),
]
for title, desc in scenarios:
    p = doc.add_paragraph()
    r = p.add_run(f'{title}: ')
    r.bold = True
    p.add_run(desc)

doc.add_page_break()

# ============================================================
# SECTION 8: LOGGING SYSTEM
# ============================================================
doc.add_heading('8. Comprehensive Logging System', level=1)

doc.add_heading('8.1 Design Philosophy: Log Everything', level=2)
doc.add_paragraph(
    'The logging system captures EVERY piece of data returned by every API call for every file '
    'on every sync cycle. There is no filtering of "irrelevant" data. If the API returns it, we store it. '
    'UNCHANGED files are logged as unchanged with their current state. The log is the single source of '
    'truth for what happened, when, and by whom.'
)

doc.add_heading('8.2 Log Entry Schema (Exhaustive)', level=2)
doc.add_paragraph('See Appendix B for the complete JSON schema. Key data captured per file per sync:')
add_table(doc,
    ['Category', 'Fields Captured', 'Source API'],
    [
        ['Sync run metadata', 'sync_id, trigger type, trigger_by, start/end timestamps, duration, GitHub run ID/URL, runner info, script version, config hash, totals per category', 'System'],
        ['Folder scan', 'folder name, Drive ID, web link, file count, subfolder count, subfolder list, total size, scan duration', 'Drive API'],
        ['Parent folder activity', 'New folders/files, created by, created when, tracked/untracked status', 'Drive Activity API'],
        ['File identity', 'name, Drive ID, mime type, extension, original filename, web view link, web content link, icon link, thumbnail link', 'Drive API'],
        ['File properties', 'size bytes, md5 checksum, sha256 (computed), created time, modified time, version, head revision ID, quota bytes used, starred, shared, trashed, description', 'Drive API'],
        ['File capabilities', 'canEdit, canComment, canShare, canCopy, canDownload, canTrash, canRename, canReadRevisions (full capabilities object)', 'Drive API'],
        ['User: last modifier', 'email, display name, Google user ID, photo link', 'Drive API + People API'],
        ['User: owner(s)', 'email, display name for each owner', 'Drive API'],
        ['User: permissions', 'permission ID, type, role, email for each permission entry', 'Drive API'],
        ['Revision history', 'All revisions since last sync: revision ID, timestamp, modifier email/name, size, md5, mime type, keep_forever flag, published flag', 'Drive API (revisions.list)'],
        ['Drive Activity', 'All actions: type (edit/rename/move/delete/comment/share/create), actor, target, timestamp, detail', 'Drive Activity API'],
        ['Previous sync state', 'Previous revision ID, md5, size, synced timestamp', 'sync_state.json'],
        ['File status', 'UNCHANGED / MODIFIED / ADDED / DELETED', 'Computed'],
        ['Sync action', 'synced / skipped / held / rejected', 'Computed'],
        ['Validation results', 'Each check: passed/failed, details, reason', 'Computed'],
        ['Content extraction', 'Method used (google_slides_api / binary_parse_python_pptx / none_unsupported_format), success/failure', 'Conversion pipeline'],
        ['Content diff (after conversion)', 'Slides modified, speaker notes changed, tables modified (cell-level), text added/removed', 'Diff engine'],
    ]
)

doc.add_heading('8.3 Log Security (Read-Only Protection)', level=2)
security_points = [
    'Log file location: logs/activity_log.json in the GitHub repository',
    'Write access: ONLY the GitHub Actions bot (via GITHUB_TOKEN) can write to the logs/ directory',
    'Branch protection rule: Require status checks to pass before merging to main; protect the logs/ path',
    'Append-only: The sync script only APPENDS new entries. It never modifies or deletes existing entries.',
    'Git history: Even if the file were tampered with, all previous versions are preserved in Git history and cannot be deleted without force-push (which is blocked by branch protection)',
    'Read access: Admin (you) and designated users can read the log via the GitHub repo',
    'The log file will grow over time. A rotation strategy can be implemented (e.g., monthly log files: activity_log_2026_02.json) to keep individual files manageable.',
]
for s in security_points:
    add_bullet(doc, s)

doc.add_heading('8.4 Activity Tracking Capabilities & Limitations', level=2)
add_table(doc,
    ['Event', 'Can We Track?', 'API', 'Latency'],
    [
        ['File edited/saved', 'YES', 'Drive API + Drive Activity API', 'Near real-time (detected at next sync)'],
        ['File created', 'YES', 'Drive API + Drive Activity API', 'Near real-time'],
        ['File deleted', 'YES', 'Drive API (absence) + Activity API', 'Near real-time'],
        ['File renamed', 'YES', 'Drive Activity API', 'Near real-time'],
        ['File moved', 'YES', 'Drive Activity API', 'Near real-time'],
        ['File commented on', 'YES', 'Drive Activity API', 'Near real-time'],
        ['Sharing/permissions changed', 'YES', 'Drive Activity API', 'Near real-time'],
        ['File viewed/opened', 'NO (requires Admin SDK)', 'Admin SDK Reports API', 'DEFERRED \u2014 needs Workspace admin'],
        ['File downloaded', 'NO (requires Admin SDK)', 'Admin SDK Reports API', 'DEFERRED'],
        ['File closed', 'NO', 'Not available in any Google API', 'Not possible'],
        ['Who has file open now', 'NO (for Office files)', 'Not available for pptx/docx', 'Not possible'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 9: GITHUB REPO
# ============================================================
doc.add_heading('9. GitHub Repository Design', level=1)

doc.add_heading('9.1 Repository Structure', level=2)
add_code_block(doc, """curriculum-kb/  (private repo)
\u251c\u2500\u2500 sources/
\u2502   \u251c\u2500\u2500 term1/                    \u2190 synced from Google Drive
\u2502   \u2502   \u251c\u2500\u2500 lessons/
\u2502   \u2502   \u251c\u2500\u2500 assessment/
\u2502   \u2502   \u2514\u2500\u2500 resources/
\u2502   \u251c\u2500\u2500 term2/                    \u2190 synced from Google Drive
\u2502   \u2502   \u251c\u2500\u2500 lessons/
\u2502   \u2502   \u251c\u2500\u2500 assessment/
\u2502   \u2502   \u2514\u2500\u2500 resources/
\u2502   \u2514\u2500\u2500 term3/                    \u2190 synced from Google Drive
\u251c\u2500\u2500 output/
\u2502   \u251c\u2500\u2500 term1_kb.json             \u2190 generated (do not edit)
\u2502   \u251c\u2500\u2500 term2_kb.json             \u2190 generated
\u2502   \u251c\u2500\u2500 term3_kb.json             \u2190 generated
\u2502   \u2514\u2500\u2500 changelog.md              \u2190 generated per build
\u251c\u2500\u2500 logs/
\u2502   \u2514\u2500\u2500 activity_log.json         \u2190 append-only, protected
\u251c\u2500\u2500 scripts/
\u2502   \u251c\u2500\u2500 sync.py                   \u2190 Drive \u2192 Git sync + validation
\u2502   \u251c\u2500\u2500 convert.py                \u2190 pptx/docx/native \u2192 JSON
\u2502   \u251c\u2500\u2500 diff_generator.py         \u2190 content-level diff + changelog
\u2502   \u2514\u2500\u2500 notify.py                 \u2190 Slack notifications
\u251c\u2500\u2500 config.yaml                       \u2190 tracked folders, Slack, schedule
\u251c\u2500\u2500 sync_state.json                   \u2190 last-known state of all files
\u251c\u2500\u2500 .github/
\u2502   \u2514\u2500\u2500 workflows/
\u2502       \u251c\u2500\u2500 sync.yml              \u2190 Drive sync workflow
\u2502       \u2514\u2500\u2500 build.yml             \u2190 KB conversion workflow
\u2514\u2500\u2500 .gitignore""")

doc.add_heading('9.2 Branch Protection', level=2)
protection = [
    'Main branch protection enabled',
    'logs/ directory: Only GitHub Actions bot can write (enforced by CODEOWNERS + required checks)',
    'output/ directory: Only generated by the conversion pipeline',
    'No force-push allowed on main',
    'All automated commits use the GitHub Actions bot identity',
]
for p_text in protection:
    add_bullet(doc, p_text)

doc.add_heading('9.3 GitHub Actions Workflows', level=2)
doc.add_paragraph('Two separate workflows:')

p = doc.add_paragraph()
r = p.add_run('sync.yml \u2014 Drive Sync: ')
r.bold = True
p.add_run('Detects changes on Drive, validates, downloads, commits to Git, notifies Slack.')

p = doc.add_paragraph()
r = p.add_run('build.yml \u2014 KB Conversion: ')
r.bold = True
p.add_run('Triggered after sync.yml commits new source files. Converts pptx/docx/native to JSON, '
          'generates changelog, commits output, notifies Slack.')

doc.add_paragraph('Separation prevents infinite loops: sync.yml triggers on schedule/manual, build.yml '
                   'triggers on source file changes. Output commits don\'t trigger either workflow.')

doc.add_page_break()

# ============================================================
# SECTION 10: CONVERSION PIPELINE
# ============================================================
doc.add_heading('10. Conversion Pipeline', level=1)

doc.add_heading('10.1 Extraction Strategy', level=2)
add_table(doc,
    ['Source Format', 'Method', 'Extracts'],
    [
        ['Google Slides (native)', 'Slides API v1', 'Slides (text, shapes, tables, speaker notes, layouts) as structured JSON'],
        ['Google Docs (native)', 'Docs API v1', 'Paragraphs, headings, tables, lists, inline objects as structured JSON'],
        ['Google Sheets (native)', 'Sheets API v4', 'Cell values, ranges, headers, merged cells as structured JSON'],
        ['PowerPoint (.pptx)', 'python-pptx', 'Slides (text, shapes, tables, speaker notes, slide order)'],
        ['Word (.docx)', 'python-docx', 'Paragraphs, headings, tables, lists, styles'],
        ['PDF', 'Not extracted', 'Stored as reference material only, logged'],
        ['Video', 'Not extracted', 'Logged with metadata only'],
    ]
)

doc.add_heading('10.2 Table Preservation', level=2)
doc.add_paragraph(
    'Tables (especially rubric matrices) are the most critical structural element. Every table '
    'is converted to a structured array format that preserves exact row/column alignment:'
)
add_code_block(doc, """{
  "title": "Game Product Rubric (50%)",
  "source": "Lesson 4 - Rewriting the Brief.pptx",
  "slide_number": 8,
  "headers": ["Criteria", "Approaching", "Meeting", "Above Expectations"],
  "rows": [
    ["Design coherence", "Basic layout...", "Clear themed...", "Polished integrated..."],
    ["Mechanic implementation", "Simple...", "Working...", "Complex, balanced..."],
    ["User experience", "Functional...", "Intuitive...", "Seamless, engaging..."]
  ],
  "row_count": 3,
  "col_count": 4
}

Guarantees: row[i][j] always corresponds to headers[j].
Merged cells: Repeated to maintain column alignment.
Empty cells: Represented as "" (empty string).
Multi-line cells: Line breaks preserved within the string.""")

doc.add_heading('10.3 Content-Level Diff', level=2)
doc.add_paragraph(
    'After conversion, the diff engine compares the newly extracted JSON against the previous '
    'version (stored from the last build). This produces a content-level diff showing exactly '
    'what changed inside each file:'
)
add_code_block(doc, """{
  "file": "Lesson 4 - Rewriting the Brief.pptx",
  "content_diff": {
    "slides_modified": [5, 8],
    "slides_added": [13],
    "slides_removed": [],
    "speaker_notes_changed": [5],
    "tables_modified": [
      {
        "slide": 8,
        "table_title": "Game Product Rubric",
        "changes": [
          {"type": "cell_changed", "row": 2, "col": 1,
           "old": "Basic layout", "new": "Clear themed layout"},
          {"type": "row_added", "row": 3,
           "content": ["User experience", "Functional...", ...]}
        ]
      }
    ],
    "text_added": ["New activity: Team Charter exercise (slide 13)"],
    "text_removed": [],
    "text_modified": [
      {"slide": 5, "old": "Form teams of 3", "new": "Form teams of 3-4"}
    ]
  }
}""")

doc.add_page_break()

# ============================================================
# SECTION 11: JSON SCHEMA
# ============================================================
doc.add_heading('11. JSON Schema Design (Backward-Compatible)', level=1)

doc.add_heading('11.1 Original Fields (Preserved)', level=2)
doc.add_paragraph('All existing fields from the current KB JSON are preserved exactly as-is:')
add_table(doc,
    ['Field', 'Type', 'Status'],
    [
        ['lesson_title', 'string', 'Preserved'],
        ['url', 'string', 'Preserved'],
        ['metadata.term_id', 'int', 'Preserved'],
        ['metadata.lesson_id', 'int', 'Preserved'],
        ['metadata.title', 'string', 'Preserved'],
        ['metadata.grade_band', 'string', 'Preserved'],
        ['metadata.core_topics', 'array<string>', 'Preserved'],
        ['metadata.endstar_tools', 'array<string>', 'Preserved'],
        ['metadata.ai_focus', 'array<string>', 'Preserved'],
        ['metadata.learning_objectives', 'array<string>', 'Preserved'],
        ['metadata.activity_type', 'string', 'Preserved'],
        ['metadata.activity_description', 'string', 'Preserved'],
        ['metadata.artifacts', 'array<string>', 'Preserved'],
        ['metadata.assessment_signals', 'array<string>', 'Preserved'],
        ['metadata.videos', 'array', 'Preserved'],
        ['metadata.resources', 'array<string>', 'Preserved'],
        ['metadata.keywords', 'array<string>', 'Preserved'],
        ['metadata.images', 'array<object>', 'Preserved (with full visual descriptions)'],
        ['description_of_activities', 'string', 'Preserved'],
        ['other_resources', 'string', 'Preserved'],
        ['videos_column', 'string', 'Preserved'],
        ['testing_scores', 'string', 'Preserved'],
        ['comments', 'string', 'Preserved'],
        ['prompts', 'string', 'Preserved'],
    ]
)

doc.add_heading('11.2 Enriched Block (New)', level=2)
add_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['enriched.key_facts', 'array<string>', 'Critical facts the chatbot must know for accurate responses'],
        ['enriched.detailed_activities', 'array<object>', 'Activity breakdown: id, title, description, slide_references'],
        ['enriched.rubrics', 'array<object>', 'Full rubric tables: title, headers[], rows[][], preserving row/col alignment'],
        ['enriched.teacher_notes', 'array<string>', 'Speaker notes from pptx, prefixed with slide reference'],
        ['enriched.assessment_framework', 'object', 'Weights, scoring methods, rubric types'],
        ['enriched.source_files', 'array<string>', 'Source files that contributed to this lesson'],
        ['enriched.last_updated', 'ISO 8601', 'When enriched data was last regenerated'],
        ['enriched.extraction_method', 'string', 'How content was extracted (API or binary parse)'],
        ['enriched.content_diff', 'object', 'What changed since previous build (cell-level detail)'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 12: NOTIFICATIONS
# ============================================================
doc.add_heading('12. Notification System (Slack)', level=1)

doc.add_heading('12.1 Activity Alerts (per sync)', level=2)
add_code_block(doc, """Slack #kb-updates:

\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501
  SYNC COMPLETE \u2014 25 Feb 2026, 00:00 UAE
\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501
  Files scanned: 146
  Changed: 4 | Added: 1 | Deleted: 0 | Held: 1

  SYNCED:
    Term 2 / Lesson 4.pptx \u2014 edited by ahmed@
    Term 2 / Lesson 7.pptx \u2014 edited by sarah@
    Term 3 / Lesson 11.pptx \u2014 edited by fatima@
    Term 1 / Assessment Guide.docx \u2014 edited by ahmed@

  ADDED:
    Term 3 / Lesson 15.pptx \u2014 new file by fatima@

  HELD (admin review needed):
    Term 2 / "Lesson 4.pptx" \u2014 potential duplicate of
      "Lesson 4 - Rewriting the Brief.pptx"

  KB build will start automatically.
\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501""")

doc.add_heading('12.2 Build Notifications', level=2)
add_code_block(doc, """Slack #kb-updates:

\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501
  KB BUILD COMPLETE \u2014 25 Feb 2026, 00:03 UAE
\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501

  Term 2 \u2014 Lesson 4:
    + 5 key facts added (Team Charter, team roles)
    + 3 activities added (Formation, Insight Wall, Rewriting)
    + 1 rubric table (3 criteria x 3 levels)
    ~ Slide 5 speaker notes updated
    ~ Rubric row 2 col 1: "Basic layout" \u2192 "Clear themed layout"

  Term 2 \u2014 Lesson 7:
    + 8 teacher notes from speaker notes
    + 5 detailed activities

  Term 3 \u2014 Lesson 11:
    + Initial enriched data (first extraction)

  Files: term1_kb.json (no changes)
         term2_kb.json (updated)
         term3_kb.json (updated)

  Full changelog: https://github.com/.../changelog.md
  cc: @admin @curriculum-team
\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501""")

doc.add_heading('12.3 Error & Warning Alerts', level=2)
add_code_block(doc, """Slack #kb-updates:

  \u26a0\ufe0f DUPLICATE DETECTED
  New: "Lesson 4.pptx" (sarah@, 3.2MB)
  Existing: "Lesson 4 - Rewriting the Brief.pptx" (ahmed@, 4.5MB)
  Match type: Same lesson number
  Action: File held. Admin review needed.

  \u274c CORRUPT FILE
  "Lesson 8.pptx" could not be parsed by python-pptx.
  Uploaded by: ahmed@ at 3:15pm
  Error: BadZipFile - file may be damaged
  Action: File rejected. Other files synced normally.

  \U0001f6a8 FILE DELETED
  "Lesson 3 - Research Methods.pptx" deleted from Term 2 by fatima@
  KB NOT updated \u2014 previous version preserved.
  Admin: Confirm deletion or restore from Drive trash.

  \U0001f4c1 PARENT FOLDER ACTIVITY
  New folder "Meeting Notes" created by sarah@
  New file "agenda.docx" added to parent by ahmed@
  These are NOT tracked. Add to config.yaml to include.""")

doc.add_heading('12.4 Distribution List', level=2)
doc.add_paragraph('Managed in config.yaml. Anyone on the list gets mentioned in Slack notifications:')
add_code_block(doc, """slack:
  webhook_url: "https://hooks.slack.com/services/xxx/yyy/zzz"
  channel: "#kb-updates"
  notify:
    - "@admin-user"
    - "@curriculum-team"
    - "@chatbot-maintainer"
  admin_only_alerts:
    - "@admin-user"   # Gets deletion, duplicate, and error alerts""")

doc.add_page_break()

# ============================================================
# SECTION 13: CONFLICT & ERROR HANDLING
# ============================================================
doc.add_heading('13. Conflict & Error Handling', level=1)

doc.add_paragraph(
    'There is no "conflict" in the traditional Git sense. Google Drive handles concurrent editing '
    'by creating sequential revisions. Two people editing the same file is normal workflow \u2014 '
    'each save creates a new revision, and the sync always downloads the latest. All revisions '
    'are logged with who made each one and when.'
)

doc.add_heading('Error Recovery Matrix', level=2)
add_table(doc,
    ['Scenario', 'System Response', 'Recovery'],
    [
        ['Corrupt file', 'File rejected, others continue, Slack alert', 'User re-uploads fixed file, next sync picks it up'],
        ['Accidental deletion', 'KB keeps previous version, Slack alert', 'Restore from Drive trash (30 days) or Git history'],
        ['Duplicate filename', 'New file held, Slack alert', 'Admin renames or deletes duplicate on Drive'],
        ['Lesson count drop', 'Entire term build held, Slack alert', 'Admin investigates and resolves'],
        ['Drive API failure', 'Sync fails, Slack alert, GitHub emails', 'Automatic retry on next scheduled sync'],
        ['GitHub Actions failure', 'Build fails, GitHub sends email', 'Admin checks Actions tab, re-triggers manually'],
        ['Slack webhook failure', 'Build succeeds silently', 'GitHub email notifications as backup; fix webhook'],
        ['OAuth token expired', 'Sync fails, Slack alert', 'Admin refreshes token (automatic if refresh token valid)'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 14: TEAM WORKFLOW
# ============================================================
doc.add_heading('14. Team Workflow (Zero Extra Steps)', level=1)

doc.add_heading('For Content Creators (Teachers, Curriculum Designers)', level=2)
doc.add_paragraph('Their workflow does NOT change at all:')
creator_steps = [
    'Open the file on Google Drive (as they always do)',
    'Edit in PowerPoint, Word, Google Slides, or Google Docs',
    'Save',
    'Done. The pipeline handles everything else automatically.',
]
for i, s in enumerate(creator_steps):
    doc.add_paragraph(f'{i+1}. {s}')

doc.add_paragraph('They will see Slack notifications about what was synced and built, but no action is required from them.')

doc.add_heading('For Admin', level=2)
admin_tasks = [
    'On-demand sync: Go to GitHub Actions tab \u2192 Sync workflow \u2192 "Run workflow" button',
    'Add new term folder: Edit config.yaml, add the Drive folder ID to tracked_folders',
    'Update Slack distribution list: Edit config.yaml slack.notify section',
    'Review held files: Check Slack alerts, resolve on Google Drive (rename/delete duplicate, confirm deletion)',
    'Monitor builds: GitHub Actions tab shows full build logs and history',
    'Review audit logs: Read logs/activity_log.json in the repo',
]
for a in admin_tasks:
    add_bullet(doc, a)

doc.add_heading('For Chatbot Maintainer', level=2)
maint_tasks = [
    'Watch Slack #kb-updates for build complete notifications',
    'Pull latest term1_kb.json, term2_kb.json, term3_kb.json from the repo',
    'Review changelog.md to understand what changed',
    'Update chatbot KB endpoint with new files',
    'Existing fields unchanged \u2014 no integration changes needed unless adopting enriched{} data',
]
for m in maint_tasks:
    add_bullet(doc, m)

doc.add_page_break()

# ============================================================
# SECTION 15: TESTING PLAN
# ============================================================
doc.add_heading('15. Testing & Simulation Plan', level=1)

doc.add_heading('Phase 1: Pipeline Validation (Admin solo)', level=2)
phase1 = [
    'Set up GitHub repo with folder structure and scripts',
    'Run sync.py locally against the real Drive folders',
    'Verify all 146 files are detected and logged correctly',
    'Verify native Google files are read via Slides/Docs/Sheets APIs',
    'Verify Office files are downloaded and parsed by python-pptx/docx',
    'Verify sync_state.json is populated correctly',
    'Verify activity_log.json captures full data for every file',
    'Test conversion pipeline: verify JSON output matches expected schema',
    'Test backward compatibility: compare output with existing Term 2 KB JSON',
    'Verify Slack notifications arrive correctly',
]
for p_text in phase1:
    add_bullet(doc, p_text)

doc.add_heading('Phase 2: Scenario Testing (Admin + 1 colleague)', level=2)
phase2 = [
    'Test normal edit: Edit a file on Drive, run sync, verify it\'s picked up',
    'Test duplicate detection: Upload a file with similar name, verify it\'s caught',
    'Test deletion detection: Delete a file, run sync, verify KB is preserved and alert sent',
    'Test corrupt file: Upload a broken pptx, verify it\'s rejected gracefully',
    'Test new subfolder: Create a new Week folder inside a term, verify it\'s detected',
    'Test new root folder: Create a folder in the parent directory, verify notification',
    'Test on-demand trigger: Manually trigger sync from GitHub Actions',
    'Test logging completeness: Review activity_log.json for exhaustive data capture',
    'Test content-level diff: Make a specific table change, verify the diff captures it at cell level',
]
for p_text in phase2:
    add_bullet(doc, p_text)

doc.add_heading('Phase 3: Guided Pilot (2-3 content creators)', level=2)
phase3 = [
    'Content creators make real edits to lesson plans on Google Drive',
    'Observe: Do Slack notifications arrive? Are they useful and clear?',
    'Verify: Are changes correctly reflected in the built KB JSON?',
    'Ask: Is anything confusing or unexpected from the creator\'s perspective?',
    'Run for 1-2 weeks before full rollout',
    'Refine notification messages based on feedback',
]
for p_text in phase3:
    add_bullet(doc, p_text)

doc.add_page_break()

# ============================================================
# SECTION 16: IMPLEMENTATION PHASES
# ============================================================
doc.add_heading('16. Implementation Phases', level=1)

add_table(doc,
    ['Phase', 'Deliverables', 'Dependencies', 'Priority'],
    [
        ['1. Google Cloud + OAuth', 'APIs enabled, OAuth flow working, token stored, Drive folder scanning verified', 'Google Cloud project (DONE)', 'Immediate'],
        ['2. GitHub Repo Setup', 'Repo created, folder structure, config.yaml, .gitignore, branch protection, secrets configured', 'Phase 1 (token for secrets)', 'Immediate'],
        ['3. Sync Pipeline', 'sync.py: change detection, full logging, validation checks, file download, Git commit', 'Phase 1 + 2', 'High'],
        ['4. Conversion Pipeline', 'convert.py: pptx/docx/native extraction, table preservation, JSON generation with enriched{} block', 'Phase 3', 'High'],
        ['5. Diff Engine', 'diff_generator.py: content-level comparison, changelog generation', 'Phase 4', 'Medium'],
        ['6. Slack Notifications', 'notify.py: sync alerts, build alerts, error alerts, parent folder monitoring', 'Phase 3', 'Medium'],
        ['7. GitHub Actions', 'sync.yml + build.yml workflows, cron schedule, manual trigger', 'Phase 3 + 4 + 5 + 6', 'High'],
        ['8. Testing', 'All three testing phases completed', 'Phase 7', 'Critical'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 17: APPENDIX A - CONFIG
# ============================================================
doc.add_heading('17. Appendix A: config.yaml Reference', level=1)
add_code_block(doc, """# Curriculum KB Maintenance Mechanism - Configuration
# This file controls the sync pipeline behavior.

# Tracked Google Drive folders
tracked_folders:
  - name: "Term 1"
    drive_folder_id: "17s13FlHGkaNPPlf3jAUY0tSza2yxHqPe"
    description: "Term 1 Content (Foundations) - January 2026 Cohort"
  - name: "Term 2"
    drive_folder_id: "16UgEwue1ROxFJyPTrowIqTQyduoNEIUb"
    description: "Term 2 Content (Accelerator) - August 2025 Cohort"
  - name: "Term 3"
    drive_folder_id: "1T6zzl0oqltIGcl8M4wAg2xy-z2HDZuxi"
    description: "Term 3 Content (Mastery) - August 2025 Cohort"

# Parent folder (for monitoring untracked activity)
parent_folder_id: "PARENT_FOLDER_ID_HERE"

# Schedule (cron expression, UTC)
schedule: "0 20 * * *"   # Midnight UAE (UTC+4 = 20:00 UTC)

# Slack configuration
slack:
  webhook_url: "https://hooks.slack.com/services/xxx/yyy/zzz"
  channel: "#kb-updates"
  notify:                  # Mentioned on all notifications
    - "@admin"
    - "@curriculum-team"
  admin_only:              # Mentioned on errors/warnings only
    - "@admin"

# Admin emails (for force operations)
admin_emails:
  - "admin@company.com"

# Conversion settings
conversion:
  extract_speaker_notes: true
  extract_images: false         # Handled by existing pipeline
  preserve_table_structure: true
  output_dir: "output/"

# File handling
file_handling:
  supported_for_kb: [".pptx", ".docx"]
  supported_native: ["application/vnd.google-apps.presentation",
                     "application/vnd.google-apps.document",
                     "application/vnd.google-apps.spreadsheet"]
  ignored_extensions: [".mp4", ".mov", ".pdf", ".zip"]
  store_but_skip_kb: [".pdf", ".xlsx"]    # Store in Git but don't extract for KB
  duplicate_similarity_threshold: 0.8

# Logging
logging:
  log_unchanged_files: true     # Log EVERY file even if unchanged
  log_file: "logs/activity_log.json"
  rotation: "monthly"           # Creates new file each month""", font_size=7)

doc.add_page_break()

# ============================================================
# SECTION 18: APPENDIX B - LOG SCHEMA
# ============================================================
doc.add_heading('18. Appendix B: Full Log Entry Schema', level=1)
doc.add_paragraph('Each sync run produces one entry containing the following structure. '
                   'Every field shown is captured from the Google APIs with no filtering.')
add_code_block(doc, """{
  "sync_run": {
    "sync_id": "sync-2026-02-25-200000",
    "trigger": "scheduled | on_demand",
    "triggered_by": "cron | admin_email@company.com",
    "started_at": "ISO 8601",
    "completed_at": "ISO 8601",
    "duration_seconds": 222,
    "github_run_id": "12345678",
    "github_run_url": "https://github.com/.../actions/runs/12345678",
    "runner_os": "ubuntu-latest",
    "script_version": "1.0.0",
    "config_hash": "sha256:...",
    "tracked_folders_count": 3,
    "total_files_scanned": 146,
    "files_unchanged": 140,
    "files_modified": 4,
    "files_added": 1,
    "files_deleted": 0,
    "files_held": 1,
    "files_rejected": 0,
    "files_synced": 5,
    "new_untracked_items": ["Meeting Notes (folder)"],
    "errors": []
  },

  "folder_scans": [ ... per folder ... ],
  "parent_folder_activity": [ ... new items in parent ... ],

  "file_entries": [
    {
      "status": "UNCHANGED | MODIFIED | ADDED | DELETED",
      "sync_action": "synced | skipped | held | rejected",
      "held_reason": null | "string describing why",

      "file": {
        "name": "string",
        "drive_id": "string",
        "mime_type": "string",
        "is_native_google": true | false,
        "native_type": null | "google_slides" | "google_doc" | "google_sheet",
        "file_extension": "string",
        "full_file_extension": "string",
        "original_filename": "string",
        "size_bytes": 0,
        "md5_checksum": "string",
        "sha256_checksum": "string (computed)",
        "created_time": "ISO 8601",
        "modified_time": "ISO 8601",
        "version": "string",
        "head_revision_id": "string",
        "web_view_link": "URL",
        "web_content_link": "URL",
        "icon_link": "URL",
        "thumbnail_link": "URL",
        "starred": false,
        "shared": true,
        "trashed": false,
        "explicitly_trashed": false,
        "description": "string",
        "quota_bytes_used": "string",
        "folder_name": "Term 2",
        "folder_path": "Curriculum Materials/Term 2/Week 2",
        "folder_drive_id": "string",

        "last_modifying_user": {
          "email": "string",
          "display_name": "string",
          "google_user_id": "string",
          "photo_link": "URL"
        },
        "owners": [{"email": "...", "display_name": "..."}],
        "permissions": [{"id": "...", "type": "...", "role": "...", "email": "..."}],
        "capabilities": { "canEdit": true, "canComment": true, ... }
      },

      "revisions_since_last_sync": [
        {
          "revision_id": "string",
          "modified_time": "ISO 8601",
          "modified_by": {"email": "...", "display_name": "..."},
          "size_bytes": 0,
          "md5_checksum": "string",
          "mime_type": "string",
          "keep_forever": false,
          "published": false
        }
      ],

      "drive_activity": [
        {
          "action_type": "edit | rename | move | delete | comment | share | create",
          "timestamp": "ISO 8601",
          "actor": {"email": "...", "display_name": "..."},
          "detail": "string (rename: old_name -> new_name, move: old_path -> new_path, etc.)"
        }
      ],

      "previous_sync_state": {
        "revision_id": "string",
        "md5_checksum": "string",
        "size_bytes": 0,
        "synced_at": "ISO 8601"
      },

      "validation": {
        "duplicate_check": {"passed": true, "details": null},
        "integrity_check": {"passed": true, "details": "parsed successfully"},
        "deletion_check": {"passed": true, "details": null}
      },

      "content_extraction": {
        "method": "google_slides_api | binary_parse_python_pptx | none_unsupported",
        "success": true,
        "error": null
      }
    }
  ]
}""", font_size=6)

doc.add_page_break()

# ============================================================
# SECTION 19: APPENDIX C - GOOGLE CLOUD SETUP
# ============================================================
doc.add_heading('19. Appendix C: Google Cloud Project Setup', level=1)

setup_steps = [
    'Google Cloud Project: poetic-dock-483707-n1',
    'APIs enabled: Drive API v3, Drive Activity API v2, Slides API v1, Docs API v1, Sheets API v4, People API',
    'Admin SDK Reports API: DEFERRED (requires Workspace admin privileges)',
    'OAuth Client: Installed application (Desktop), Client ID: 509243096178-*.apps.googleusercontent.com',
    'Token: Stored locally as token.json (for development) and as GitHub Actions secret GOOGLE_TOKEN_JSON (for production)',
    'Scopes: drive.readonly, drive.activity.readonly, presentations.readonly, documents.readonly, spreadsheets.readonly, directory.readonly',
]
for s in setup_steps:
    add_bullet(doc, s)

# ============================================================
# END
# ============================================================
doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\u2014\u2014\u2014 End of Document \u2014\u2014\u2014')
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r.font.size = Pt(9)

# Save
output = 'KB_Maintenance_Mechanism_Pipeline_v2.docx'
doc.save(output)
print(f'Document saved: {output}')
